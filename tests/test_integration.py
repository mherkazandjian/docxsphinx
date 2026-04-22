"""Integration tests — visitor + python-docx on parsed source strings.

Unit tests in ``test_visitor_units.py`` construct doctrees by hand to
isolate the visitor; e2e tests in ``test_build_examples.py`` shell out to
``sphinx-build``. This module sits between: parse real RST or MD through
the public docutils / myst-parser programmatic API, run the
``DocxTranslator`` over the resulting doctree, and inspect the python-docx
``Document`` in memory.

The goal is to catch regressions in how *parser output* combines with the
*writer* — neither tier alone covers this seam.
"""
from __future__ import annotations

from collections.abc import Callable

import pytest
from docx.document import Document as DocumentType

pytestmark = pytest.mark.integration


def _list_styles(doc: DocumentType) -> list[str]:
    return [(p.style.name if p.style else '-') for p in doc.paragraphs]


def _list_texts(doc: DocumentType) -> list[str]:
    return [p.text for p in doc.paragraphs]


# ---------------------------------------------------------------------------
# RST pipeline
# ---------------------------------------------------------------------------

def test_rst_headings_and_paragraphs(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    src = """Top heading
===========

Introductory paragraph under the top heading.

Second section
--------------

A paragraph under the second section.

Third section
-------------

A paragraph under the third section.
"""
    doc = translator_factory(src, 'rst')
    texts = _list_texts(doc)
    styles_by_text = {p.text: (p.style.name if p.style else '-') for p in doc.paragraphs}

    # Each section title should render via some heading-family style (docutils
    # may promote the outermost ``===`` section to the document Title, so we
    # accept either ``Title`` or ``Heading N`` for any of the three).
    for title in ('Top heading', 'Second section', 'Third section'):
        assert title in texts, (title, texts)
        style = styles_by_text[title]
        assert style == 'Title' or style.startswith('Heading'), (title, style)


def test_rst_strong_and_emphasis(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    doc = translator_factory('A line with **bold** and *italic* spans.\n', 'rst')
    bold_texts = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
    italic_texts = [r.text for p in doc.paragraphs for r in p.runs if r.italic]
    assert any('bold' in t for t in bold_texts), bold_texts
    assert any('italic' in t for t in italic_texts), italic_texts


def test_rst_bullet_list_uses_list_bullet_style(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    doc = translator_factory('* alpha\n* beta\n* gamma\n', 'rst')
    list_styles = [s for s in _list_styles(doc) if 'List' in s]
    assert list_styles.count('List Bullet') >= 3, list_styles


# ---------------------------------------------------------------------------
# MD pipeline — confirm MyST feeds the visitor the same kinds of nodes
# ---------------------------------------------------------------------------

def test_md_headings_and_paragraphs(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    src = """# Top heading

Introductory paragraph under the top heading.

## Second section

A paragraph under the second section.

## Third section

A paragraph under the third section.
"""
    doc = translator_factory(src, 'md')
    texts = _list_texts(doc)
    styles_by_text = {p.text: (p.style.name if p.style else '-') for p in doc.paragraphs}

    for title in ('Top heading', 'Second section', 'Third section'):
        assert title in texts, (title, texts)
        style = styles_by_text[title]
        assert style == 'Title' or style.startswith('Heading'), (title, style)


def test_md_strong_and_emphasis_parity_with_rst(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """Both parsers must produce ``strong`` / ``emphasis`` nodes that the visitor flags identically."""
    doc = translator_factory('A line with **bold** and *italic* spans.\n', 'md')
    bold_texts = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
    italic_texts = [r.text for p in doc.paragraphs for r in p.runs if r.italic]
    assert any('bold' in t for t in bold_texts), bold_texts
    assert any('italic' in t for t in italic_texts), italic_texts


def test_md_bullet_list(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    doc = translator_factory('- alpha\n- beta\n- gamma\n', 'md')
    texts = _list_texts(doc)
    for item in ('alpha', 'beta', 'gamma'):
        assert item in ' '.join(texts), (item, texts)


# ---------------------------------------------------------------------------
# MD tasklist — full pipeline, not a hand-built doctree
# ---------------------------------------------------------------------------

def test_md_tasklist_end_to_end(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """Real MyST → doctree → visitor → Document: checkboxes appear as ☐/☒ glyphs."""
    src = """# checklist

- [x] done thing
- [ ] pending thing
- plain bullet
"""
    doc = translator_factory(src, 'md')
    list_items = [p.text for p in doc.paragraphs if 'List' in (p.style.name or '')]
    assert list_items == ['☒ done thing', '☐ pending thing', 'plain bullet'], list_items


# ---------------------------------------------------------------------------
# Cross-tier parity — the same logical content from both parsers should
# produce an observationally similar Document
# ---------------------------------------------------------------------------

def test_md_admonition_renders_as_table(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """Real MyST ``:::{note}`` → 1x1 table with bold ``Note`` + body."""
    src = """# doc title

:::{note}
Short note body.
:::
"""
    doc = translator_factory(src, 'md')
    assert len(doc.tables) == 1, f'expected one admonition table; got {len(doc.tables)}'
    cell = doc.tables[0].rows[0].cells[0]
    texts = [p.text for p in cell.paragraphs if p.text.strip()]
    assert texts[0] == 'Note', texts
    assert any('Short note body' in t for t in texts), texts


def test_rst_admonition_renders_as_table(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """docutils ``.. warning::`` → 1x1 table with bold ``Warning`` + body."""
    src = """
.. warning::

   Careful with the hot coffee.
"""
    doc = translator_factory(src, 'rst')
    assert len(doc.tables) == 1
    cell = doc.tables[0].rows[0].cells[0]
    texts = [p.text for p in cell.paragraphs if p.text.strip()]
    assert texts[0] == 'Warning', texts
    assert any('hot coffee' in t for t in texts), texts


def test_rst_sub_and_superscript_end_to_end(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """Real RST with ``:sub:`` / ``:sup:`` roles → runs with ``font.subscript/superscript``."""
    src = 'E = mc\\ :sup:`2`\\  and H\\ :sub:`2`\\ O are classics.\n'
    doc = translator_factory(src, 'rst')
    runs = [r for p in doc.paragraphs for r in p.runs]
    sup = [r.text for r in runs if r.font.superscript]
    sub = [r.text for r in runs if r.font.subscript]
    assert '2' in sup, sup
    assert '2' in sub, sub


def test_rst_definition_list_end_to_end(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """RST deflist → 2-col table with bold term and paragraph body."""
    src = """
alpha
   first thing

beta
   second thing
"""
    doc = translator_factory(src, 'rst')
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.rows[0].cells[0].paragraphs[0].text == 'alpha'
    assert 'first thing' in table.rows[0].cells[1].paragraphs[0].text


def test_md_deflist_end_to_end(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """MyST ``deflist`` extension → 2-col table; multi-``:`` definitions stack as paragraphs."""
    src = """
Term A
: first def

Term B
: alpha
: beta
"""
    doc = translator_factory(src, 'md')
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert table.rows[1].cells[0].paragraphs[0].text == 'Term B'
    b_paragraphs = [
        p.text for p in table.rows[1].cells[1].paragraphs if p.text.strip()
    ]
    assert b_paragraphs == ['alpha', 'beta'], b_paragraphs


def test_rst_option_list_end_to_end(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """RST option list → 2-col table with canonical option-signature strings."""
    src = """
-h, --help          Print this help message.

-o FILE, --out=FILE
                    Write output to FILE.
"""
    doc = translator_factory(src, 'rst')
    assert len(doc.tables) == 1
    rows = doc.tables[0].rows
    sigs = [r.cells[0].paragraphs[0].text for r in rows]
    assert sigs == ['-h, --help', '-o FILE, --out=FILE'], sigs


def test_rst_math_renders_as_omml(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """Full RST pipeline: ``:math:`` role → inline ``<m:oMath>``;
    ``.. math::`` directive → ``<m:oMathPara>`` in its own paragraph."""
    src = r"""
Mass-energy: :math:`E = mc^2`.

.. math::

   \sum_{i=1}^n i = \frac{n(n+1)}{2}
"""
    doc = translator_factory(src, 'rst')
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    all_oMath = sum(
        len(p._p.findall(f'.//{{{M}}}oMath')) for p in doc.paragraphs
    )
    all_oMathPara = sum(
        len(p._p.findall(f'.//{{{M}}}oMathPara')) for p in doc.paragraphs
    )
    assert all_oMath >= 2, all_oMath  # one inline + one inside oMathPara
    assert all_oMathPara == 1, all_oMathPara


def test_rst_footnotes_roundtrip_via_footnotes_part(
    translator_factory: Callable[[str, str], DocumentType],
    tmp_path,
) -> None:
    """Full RST pipeline: footnotes produce a ``word/footnotes.xml`` part
    with each body, and inline footnote references carry the matching
    OOXML ids so Word pairs them up."""
    import zipfile

    from docx.oxml.ns import qn

    src = """
Main text with a footnote [1]_ and another [2]_ here.

.. [1] First footnote body.

.. [2] Second footnote body.
"""
    doc = translator_factory(src, 'rst')
    out = tmp_path / 'fn.docx'
    doc.save(str(out))

    with zipfile.ZipFile(out) as zf:
        assert 'word/footnotes.xml' in zf.namelist()
        fn_xml = zf.read('word/footnotes.xml').decode()

    # User footnotes on top of the two standard separators.
    assert fn_xml.count('<w:footnote ') == 4, fn_xml[:400]
    assert 'First footnote body' in fn_xml
    assert 'Second footnote body' in fn_xml

    # Inline references carry ids 1 and 2 — matching the two user footnote bodies.
    ref_ids = []
    for p in doc.paragraphs:
        for ref in p._p.findall('.//' + qn('w:footnoteReference')):
            ref_ids.append(ref.get(qn('w:id')))
    assert sorted(ref_ids) == ['1', '2'], ref_ids


def test_rst_figure_caption_renders_with_seq_field(
    translator_factory: Callable[[str, str], DocumentType],
    fake_builder,
    tmp_path,
) -> None:
    """Full RST pipeline: two figures → two captions, each with a SEQ field
    and a bookmark matching the figure id."""
    import shutil
    from pathlib import Path

    from docx.oxml.ns import qn

    src_png = (
        Path(__file__).resolve().parent.parent
        / 'examples' / 'sample_1' / 'source' / 'image1.png'
    )
    shutil.copy(src_png, tmp_path / 'image1.png')
    fake_builder.env.srcdir = str(tmp_path)

    src = """
.. _fig-cats:

.. figure:: image1.png

   A pair of cats.

.. _fig-dogs:

.. figure:: image1.png

   Dogs playing.
"""
    doc = translator_factory(src, 'rst')

    caption_paragraphs = [
        p for p in doc.paragraphs if p.style and p.style.name == 'Caption'
    ]
    assert len(caption_paragraphs) == 2

    # Each caption has exactly one SEQ field.
    for p, expected_num in zip(caption_paragraphs, ['1', '2'], strict=True):
        flds = p._p.findall(qn('w:fldSimple'))
        assert len(flds) == 1
        placeholder_text = ''.join(t.text for t in flds[0].findall('.//' + qn('w:t')))
        assert placeholder_text == expected_num

    # Bookmarks for each figure id end up in the corresponding caption.
    bookmark_names = []
    for p in caption_paragraphs:
        for bm in p._p.findall(qn('w:bookmarkStart')):
            bookmark_names.append(bm.get(qn('w:name')))
    assert 'fig-cats' in bookmark_names
    assert 'fig-dogs' in bookmark_names


def test_md_fenced_python_code_is_highlighted(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """MyST fenced ``python`` code block → one code paragraph containing
    multiple coloured monospace runs."""
    src = """# code

```python
def add(a, b):
    return a + b
```
"""
    doc = translator_factory(src, 'md')
    code_paragraphs = [
        p for p in doc.paragraphs
        if any(r.font.name == 'Consolas' for r in p.runs)
    ]
    assert len(code_paragraphs) == 1
    runs = code_paragraphs[0].runs
    colours = {
        str(r.font.color.rgb)
        for r in runs
        if r.font.color is not None and r.font.color.rgb is not None
    }
    assert len(colours) >= 2, colours


def test_md_internal_anchor_link_resolves_to_bookmark(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """MyST ``[text](#section-id)`` + ``(section-id)=`` produces a
    clickable internal hyperlink whose ``w:anchor`` matches a bookmark
    emitted at the target heading."""
    from docx.oxml.ns import qn

    src = """# doc

See [below](#target-section) for details.

(target-section)=
## Target section

Body.
"""
    doc = translator_factory(src, 'md')

    # Collect all internal-anchor hyperlinks and all bookmark names.
    anchors = []
    bookmark_names = []
    for p in doc.paragraphs:
        for link in p._p.findall(qn('w:hyperlink')):
            a = link.get(qn('w:anchor'))
            if a:
                anchors.append(a)
        for bm in p._p.findall(qn('w:bookmarkStart')):
            bookmark_names.append(bm.get(qn('w:name')))

    assert 'target-section' in anchors, anchors
    assert 'target-section' in bookmark_names, bookmark_names


def test_md_external_links_render_as_hyperlinks(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """Real MyST markdown `[text](url)` → w:hyperlink w/ correct target."""
    from docx.oxml.ns import qn

    src = """# links

See [Sphinx](https://www.sphinx-doc.org/) and [PyPI](https://pypi.org/).
"""
    doc = translator_factory(src, 'md')
    # Find the paragraph containing the links (skip the heading).
    link_paragraphs = [
        p for p in doc.paragraphs if p._p.findall(qn('w:hyperlink'))
    ]
    assert link_paragraphs, [p.text for p in doc.paragraphs]
    links = link_paragraphs[0]._p.findall(qn('w:hyperlink'))
    assert len(links) == 2

    link_texts_and_targets = []
    for link in links:
        text = ''.join(t.text for t in link.findall('.//' + qn('w:t')))
        r_id = link.get(qn('r:id'))
        target = link_paragraphs[0].part.rels[r_id].target_ref
        link_texts_and_targets.append((text, target))
    assert link_texts_and_targets == [
        ('Sphinx', 'https://www.sphinx-doc.org/'),
        ('PyPI', 'https://pypi.org/'),
    ], link_texts_and_targets


def test_rst_rowspan_produces_vmerge_xml(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    """RST grid table with a rowspan → top cell gets ``w:vMerge val="restart"``,
    continuation cell gets bare ``w:vMerge``. Rows not spanned remain untouched."""
    src = """
+------+------+------+
| A    | B    | C    |
+      +------+------+
|      | D    | E    |
+------+------+------+
| F    | G    | H    |
+------+------+------+
"""
    doc = translator_factory(src, 'rst')
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert len(tbl.rows) == 3

    from docx.oxml.ns import qn

    def vmerge_layout(table):
        """Return [[val_or_None, ...per tc in row], ...per row] from raw XML.
        ``val`` is ``'restart'`` for the merge top, ``'continue'`` for the
        subsequent cells, and ``None`` for cells without any ``w:vMerge``.
        ``cells[]`` can't be used directly because python-docx collapses
        the merged-cell slots to a single wrapper."""
        out = []
        for tr in table._tbl.findall(qn('w:tr')):
            row_out = []
            for tc in tr.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                v = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
                if v is None:
                    row_out.append(None)
                else:
                    row_out.append(v.get(qn('w:val')) or 'continue')
            out.append(row_out)
        return out

    layout = vmerge_layout(tbl)
    # First row, 3 cells (A B C): A has restart vMerge, B and C have none.
    assert layout[0] == ['restart', None, None], layout
    # Second row, 3 cells (continuation of A, D, E): first is 'continue', others None.
    assert layout[1] == ['continue', None, None], layout
    # Third row, 3 independent cells: no vMerge anywhere.
    assert layout[2] == [None, None, None], layout


def test_rst_image_sizing_end_to_end(
    translator_factory: Callable[[str, str], DocumentType],
    fake_builder,
    tmp_path,
) -> None:
    """``.. image:: foo.png :width: 2in`` → inline shape whose width is 2in in EMU."""
    import shutil
    from pathlib import Path

    # Stage the sample image at the fake builder's srcdir so visit_image can find it.
    src_png = (
        Path(__file__).resolve().parent.parent
        / 'examples' / 'sample_1' / 'source' / 'image1.png'
    )
    shutil.copy(src_png, tmp_path / 'image1.png')
    fake_builder.env.srcdir = str(tmp_path)

    src = """
.. image:: image1.png
   :width: 2in
   :alt: two inch image
"""
    doc = translator_factory(src, 'rst')
    assert len(doc.inline_shapes) == 1
    shape = doc.inline_shapes[0]
    assert shape.width == 2 * 914400, shape.width  # 2in in EMU
    assert shape._inline.docPr.get('descr') == 'two inch image'


def test_rst_and_md_produce_same_paragraph_count_for_equivalent_source(
    translator_factory: Callable[[str, str], DocumentType],
) -> None:
    rst_doc = translator_factory(
        "Title\n=====\n\nA paragraph.\n\n* alpha\n* beta\n",
        'rst',
    )
    md_doc = translator_factory(
        "# Title\n\nA paragraph.\n\n- alpha\n- beta\n",
        'md',
    )
    assert len(rst_doc.paragraphs) == len(md_doc.paragraphs), (
        [p.text for p in rst_doc.paragraphs],
        [p.text for p in md_doc.paragraphs],
    )
