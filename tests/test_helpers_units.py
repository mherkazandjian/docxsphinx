"""Direct unit tests for ``src/docxsphinx/_docx_helpers``.

Each helper is exercised in isolation — no visitor, no docutils — so a
regression in the helpers surfaces at this tier rather than through a
downstream writer test. Paired with ``test_visitor_units`` this gives
two independent angles of coverage on the low-level OOXML emission.
"""
from __future__ import annotations

import io
import zipfile

import pytest
from docx import Document
from docx.oxml.ns import qn

from docxsphinx._docx_helpers import (
    add_bookmark,
    add_footnote,
    add_footnote_reference,
    add_hyperlink,
    add_internal_hyperlink,
    add_seq_field,
    ensure_footnotes_part,
)

pytestmark = pytest.mark.unit


def test_add_hyperlink_plain_text_registers_external_relationship():
    doc = Document()
    p = doc.add_paragraph('Visit: ')
    add_hyperlink(p, 'https://example.com/', 'Example')

    links = p._p.findall(qn('w:hyperlink'))
    assert len(links) == 1
    r_id = links[0].get(qn('r:id'))
    assert r_id, 'hyperlink missing r:id'
    rels = p.part.rels
    assert rels[r_id].target_ref == 'https://example.com/'
    assert rels[r_id].is_external is True

    text = ''.join(t.text for t in links[0].findall('.//' + qn('w:t')))
    assert text == 'Example'


def test_add_hyperlink_styled_chunks_emits_multiple_runs():
    """Passing a list of ``(text, bold, italic, strike)`` chunks builds one
    ``<w:r>`` per chunk with the requested formatting inside the hyperlink."""
    doc = Document()
    p = doc.add_paragraph()
    chunks = [
        ('plain ', False, False, False),
        ('bold', True, False, False),
        (' and ', False, False, False),
        ('italic', False, True, False),
        (' and ', False, False, False),
        ('struck', False, False, True),
    ]
    add_hyperlink(p, 'https://example.com/', chunks)

    link = p._p.findall(qn('w:hyperlink'))[0]
    runs = link.findall(qn('w:r'))
    assert len(runs) == 6, len(runs)

    def rpr_flags(run):
        rPr = run.find(qn('w:rPr'))
        if rPr is None:
            return (False, False, False)
        return (
            rPr.find(qn('w:b')) is not None,
            rPr.find(qn('w:i')) is not None,
            rPr.find(qn('w:strike')) is not None,
        )

    # (bold, italic, strike) per run
    assert [rpr_flags(r) for r in runs] == [
        (False, False, False),
        (True, False, False),
        (False, False, False),
        (False, True, False),
        (False, False, False),
        (False, False, True),
    ]


def test_add_internal_hyperlink_uses_anchor_not_rid():
    doc = Document()
    p = doc.add_paragraph()
    add_internal_hyperlink(p, 'my-section', 'jump here')

    link = p._p.findall(qn('w:hyperlink'))[0]
    assert link.get(qn('w:anchor')) == 'my-section'
    assert link.get(qn('r:id')) is None


def test_add_bookmark_emits_start_and_end_with_matching_ids():
    doc = Document()
    p = doc.add_paragraph('destination')
    add_bookmark(p, 'anchor-name', 42)

    starts = p._p.findall(qn('w:bookmarkStart'))
    ends = p._p.findall(qn('w:bookmarkEnd'))
    assert len(starts) == 1 and len(ends) == 1
    assert starts[0].get(qn('w:id')) == '42'
    assert ends[0].get(qn('w:id')) == '42'
    assert starts[0].get(qn('w:name')) == 'anchor-name'


def test_add_seq_field_instr_names_the_sequence():
    doc = Document()
    p = doc.add_paragraph()
    add_seq_field(p, 'Figure', initial_value=3)

    fld = p._p.find(qn('w:fldSimple'))
    assert fld is not None
    assert 'SEQ Figure' in fld.get(qn('w:instr'))
    # Placeholder numeric value rendered inline until Word refreshes fields.
    placeholder = ''.join(t.text for t in fld.findall('.//' + qn('w:t')))
    assert placeholder == '3'


def test_ensure_footnotes_part_is_idempotent():
    """Calling ``ensure_footnotes_part`` twice returns the same element
    (same id) — the part is created on first call, reused on second."""
    doc = Document()
    first = ensure_footnotes_part(doc)
    second = ensure_footnotes_part(doc)
    assert first is second

    # And the part is registered exactly once in the relationships.
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    fn_rels = [
        rel for rel in doc.part.rels.values() if rel.reltype == RT.FOOTNOTES
    ]
    assert len(fn_rels) == 1


def test_add_footnote_flat_text_writes_single_body_run():
    """Plain-string body produces a single text run after the marker."""
    doc = Document()
    root = ensure_footnotes_part(doc)
    add_footnote(root, 1, 'hello world')

    buf = io.BytesIO()
    doc.save(buf)
    with zipfile.ZipFile(buf) as zf:
        xml = zf.read('word/footnotes.xml').decode()
    assert 'hello world' in xml


def test_add_footnote_styled_chunks_emits_bold_run():
    """List-of-chunks body produces per-run formatting inside the footnote."""
    doc = Document()
    root = ensure_footnotes_part(doc)
    add_footnote(root, 1, [
        ('See ', False, False, False),
        ('bold', True, False, False),
        (' too.', False, False, False),
    ])

    buf = io.BytesIO()
    doc.save(buf)
    with zipfile.ZipFile(buf) as zf:
        xml = zf.read('word/footnotes.xml').decode()
    assert '<w:b/>' in xml or '<w:b />' in xml, xml
    assert 'See' in xml and 'bold' in xml and 'too' in xml


def test_add_footnote_reference_emits_superscripted_reference_run():
    doc = Document()
    p = doc.add_paragraph('body text')
    add_footnote_reference(p, 7)

    refs = p._p.findall('.//' + qn('w:footnoteReference'))
    assert len(refs) == 1
    assert refs[0].get(qn('w:id')) == '7'

    # Enclosing run has vertAlign=superscript.
    run = refs[0].getparent()
    rPr = run.find(qn('w:rPr'))
    vert = rPr.find(qn('w:vertAlign'))
    assert vert is not None
    assert vert.get(qn('w:val')) == 'superscript'
