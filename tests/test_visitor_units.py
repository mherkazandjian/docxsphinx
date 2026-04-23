"""Fast unit tests of ``DocxTranslator``.

These tests walk a hand-built (or quickly-parsed) doctree against the
translator without shelling out to ``sphinx-build``. They're the home for
per-feature tests added during Phase 2. Much faster feedback than the
end-to-end tests in ``test_build_examples.py``.
"""
from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

import pytest
from docutils import nodes
from docutils.core import publish_doctree
from docutils.frontend import get_default_settings
from docutils.parsers.rst import Parser as _RstParser
from docutils.utils import new_document
from docx import Document

from docxsphinx.writer import DocxTranslator

pytestmark = pytest.mark.unit


def _blank_document() -> nodes.document:
    """Return a fresh docutils ``document`` with a functional reporter."""
    settings = get_default_settings(_RstParser)
    return new_document('<test>', settings)


def _walk(rst: str, fake_builder: SimpleNamespace):
    """Parse ``rst``, walk it through a translator, return the resulting doc."""
    doctree = publish_doctree(rst)
    container = Document()
    translator = DocxTranslator(doctree, fake_builder, container)
    doctree.walkabout(translator)
    return container


def test_translator_imports_cleanly() -> None:
    """Smoke: the visitor class imports and has the expected surface."""
    assert hasattr(DocxTranslator, 'visit_paragraph')
    assert hasattr(DocxTranslator, 'visit_section')
    assert hasattr(DocxTranslator, 'add_text')


def test_simple_paragraph_emits_text(fake_builder: SimpleNamespace) -> None:
    doc = _walk("hello world\n", fake_builder)
    texts = [p.text for p in doc.paragraphs]
    assert any('hello world' in t for t in texts), texts


def test_bold_run_is_bold(fake_builder: SimpleNamespace) -> None:
    """A **bold** span should produce a run with ``bold=True``."""
    doc = _walk("this is **emphatic** text\n", fake_builder)
    bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
    assert any('emphatic' in r.text for r in bold_runs), [
        (r.text, r.bold) for p in doc.paragraphs for r in p.runs
    ]


def test_bullet_list_creates_paragraphs(fake_builder: SimpleNamespace) -> None:
    rst = "* alpha\n* beta\n* gamma\n"
    doc = _walk(rst, fake_builder)
    all_text = ' '.join(p.text for p in doc.paragraphs)
    for item in ('alpha', 'beta', 'gamma'):
        assert item in all_text, (item, all_text)
    list_styles = [
        (p.style.name if p.style else '-') for p in doc.paragraphs
        if 'List' in (p.style.name or '')
    ]
    assert list_styles == ['List Bullet', 'List Bullet', 'List Bullet'], list_styles


def test_enumerated_list_uses_list_number_style(fake_builder: SimpleNamespace) -> None:
    """RST numbered list (``1. foo\\n2. bar``) must render with
    ``List Number`` style rather than ``List Bullet``."""
    doc = _walk("1. first\n2. second\n3. third\n", fake_builder)
    list_styles = [
        (p.style.name if p.style else '-') for p in doc.paragraphs
        if 'List' in (p.style.name or '')
    ]
    assert list_styles == ['List Number', 'List Number', 'List Number'], list_styles


def test_mixed_nesting_picks_style_by_innermost_list(
    fake_builder: SimpleNamespace,
) -> None:
    """Numbered list containing a bulleted sub-list: the outer items use
    ``List Number``, the inner use ``List Bullet 2`` (level-2 depth suffix)."""
    rst = (
        "1. step one\n"
        "\n"
        "   * sub a\n"
        "   * sub b\n"
        "\n"
        "2. step two\n"
    )
    doc = _walk(rst, fake_builder)
    list_styles = [
        (p.style.name if p.style else '-') for p in doc.paragraphs
        if 'List' in (p.style.name or '')
    ]
    # Exact shape is picky about the level-bump timing — just assert the
    # mix of families is present at the right levels.
    assert 'List Number' in list_styles
    assert any('List Bullet 2' in s for s in list_styles), list_styles


def _build_tasklist_doctree(checked: bool, task_text: str) -> nodes.document:
    """Hand-build the shape of doctree MyST emits for a single GFM task item."""
    doctree = _blank_document()
    blist = nodes.bullet_list(bullet='-', classes=['contains-task-list'])
    item = nodes.list_item(classes=['task-list-item'])
    para = nodes.paragraph()
    html = (
        '<input class="task-list-item-checkbox" checked="checked" '
        'disabled="disabled" type="checkbox">'
        if checked
        else '<input class="task-list-item-checkbox" '
             'disabled="disabled" type="checkbox">'
    )
    para += nodes.raw('', html, format='html')
    para += nodes.Text(f' {task_text}')
    item += para
    blist += item
    doctree += blist
    return doctree


def _walk_doctree(doctree, fake_builder: SimpleNamespace):
    container = Document()
    translator = DocxTranslator(doctree, fake_builder, container)
    doctree.walkabout(translator)
    return container


def test_tasklist_unchecked_renders_empty_box(fake_builder: SimpleNamespace) -> None:
    """`- [ ] walk the dog` should produce a List Bullet paragraph starting with ☐."""
    doctree = _build_tasklist_doctree(checked=False, task_text='walk the dog')
    doc = _walk_doctree(doctree, fake_builder)
    list_items = [p for p in doc.paragraphs if 'List' in (p.style.name or '')]
    assert list_items, [p.text for p in doc.paragraphs]
    assert list_items[0].text == '☐ walk the dog', list_items[0].text


def test_tasklist_checked_renders_boxed_x(fake_builder: SimpleNamespace) -> None:
    """`- [x] make coffee` should produce a List Bullet paragraph starting with ☒."""
    doctree = _build_tasklist_doctree(checked=True, task_text='make coffee')
    doc = _walk_doctree(doctree, fake_builder)
    list_items = [p for p in doc.paragraphs if 'List' in (p.style.name or '')]
    assert list_items, [p.text for p in doc.paragraphs]
    assert list_items[0].text == '☒ make coffee', list_items[0].text


def test_tasklist_prefix_does_not_leak_to_following_plain_item(
    fake_builder: SimpleNamespace,
) -> None:
    """A plain bullet following a task item must not inherit ☐/☒."""
    doctree = _blank_document()
    blist = nodes.bullet_list(bullet='-')

    task = nodes.list_item(classes=['task-list-item'])
    task_para = nodes.paragraph()
    task_para += nodes.raw(
        '',
        '<input class="task-list-item-checkbox" disabled="disabled" type="checkbox">',
        format='html',
    )
    task_para += nodes.Text(' a task')
    task += task_para
    blist += task

    plain = nodes.list_item()
    plain_para = nodes.paragraph()
    plain_para += nodes.Text('plain bullet')
    plain += plain_para
    blist += plain

    doctree += blist

    doc = _walk_doctree(doctree, fake_builder)
    list_items = [p.text for p in doc.paragraphs if 'List' in (p.style.name or '')]
    assert list_items == ['☐ a task', 'plain bullet'], list_items


def test_note_admonition_renders_as_single_cell_table(
    fake_builder: SimpleNamespace,
) -> None:
    """A ``.. note::`` should produce a 1x1 table whose cell starts with a
    bold ``Note`` title paragraph and contains the body text."""
    doctree = _blank_document()
    note = nodes.note()
    note += nodes.paragraph('', 'Body of a note.')
    doctree += note

    doc = _walk_doctree(doctree, fake_builder)

    assert len(doc.tables) == 1, f'expected exactly one table; got {len(doc.tables)}'
    cell = doc.tables[0].rows[0].cells[0]
    texts = [p.text for p in cell.paragraphs if p.text.strip()]
    assert texts[0] == 'Note', texts
    assert 'Body of a note.' in ' '.join(texts), texts

    # The ``Note`` title run must be bold.
    title_runs = doc.tables[0].rows[0].cells[0].paragraphs[0].runs
    assert any(r.bold for r in title_runs), [(r.text, r.bold) for r in title_runs]


def test_warning_admonition_uses_warning_label(
    fake_builder: SimpleNamespace,
) -> None:
    doctree = _blank_document()
    warn = nodes.warning()
    warn += nodes.paragraph('', 'Danger Will Robinson.')
    doctree += warn

    doc = _walk_doctree(doctree, fake_builder)
    cell = doc.tables[0].rows[0].cells[0]
    assert cell.paragraphs[0].text == 'Warning', cell.paragraphs[0].text


def test_subscript_run_has_subscript_font(fake_builder: SimpleNamespace) -> None:
    """``H\\ :sub:`2`\\ O`` must produce a run where ``font.subscript is True``."""
    doc = _walk('H\\ :sub:`2`\\ O is water.\n', fake_builder)
    runs = [r for p in doc.paragraphs for r in p.runs]
    sub_runs = [r for r in runs if r.font.subscript]
    assert sub_runs, [(r.text, r.font.subscript) for r in runs]
    assert any(r.text == '2' for r in sub_runs), [r.text for r in sub_runs]
    # The non-subscript H and O text runs must NOT be subscripted.
    non_sub = [r for r in runs if not r.font.subscript]
    assert any('water' in r.text for r in non_sub), [r.text for r in non_sub]


def test_superscript_run_has_superscript_font(fake_builder: SimpleNamespace) -> None:
    """``E = mc\\ :sup:`2`\\ `` must produce a run with ``font.superscript is True``."""
    doc = _walk('E = mc\\ :sup:`2`\\  is Einstein.\n', fake_builder)
    runs = [r for p in doc.paragraphs for r in p.runs]
    sup_runs = [r for r in runs if r.font.superscript]
    assert sup_runs, [(r.text, r.font.superscript) for r in runs]
    assert any(r.text == '2' for r in sup_runs), [r.text for r in sup_runs]


def test_sub_and_sup_flags_clear_on_depart(fake_builder: SimpleNamespace) -> None:
    """Text after a ``:sub:`` / ``:sup:`` role must not inherit the flag."""
    doc = _walk(
        'before :sub:`inner` after :sup:`up` done.\n',
        fake_builder,
    )
    runs = [r for p in doc.paragraphs for r in p.runs]
    # Runs containing 'after' and 'done' must be plain.
    for target in ('after', 'done'):
        matches = [r for r in runs if target in r.text]
        assert matches, (target, [r.text for r in runs])
        for r in matches:
            assert not r.font.subscript, (target, r.text)
            assert not r.font.superscript, (target, r.text)


def test_length_to_emu_parses_common_units() -> None:
    """The CSS-length helper must cover px/cm/mm/in/pt + unitless + reject %."""
    from docx.shared import Cm, Emu, Inches, Mm, Pt

    from docxsphinx.writer import _length_to_emu

    assert _length_to_emu('300px') == Emu(300 * 9525)
    assert _length_to_emu('300') == Emu(300 * 9525)   # unitless == px
    assert _length_to_emu('5cm') == Cm(5)
    assert _length_to_emu('20mm') == Mm(20)
    assert _length_to_emu('2in') == Inches(2)
    assert _length_to_emu('72pt') == Pt(72)
    assert _length_to_emu('50%') is None              # percentages deferred
    assert _length_to_emu('') is None
    assert _length_to_emu(None) is None
    assert _length_to_emu('garbage') is None


def test_image_with_width_sets_inline_shape_width(
    fake_builder: SimpleNamespace, tmp_path,
) -> None:
    """An image node with ``width='300px'`` produces an inline shape whose
    width matches the parsed value."""
    import shutil

    # Copy the existing sample image into a tmpdir used as the builder srcdir.
    src_png = (
        Path(__file__).resolve().parent.parent / 'examples' / 'sample_1' / 'source' / 'image1.png'
    )
    shutil.copy(src_png, tmp_path / 'image1.png')
    fake_builder.env.srcdir = str(tmp_path)

    doctree = _blank_document()
    image = nodes.image(uri='image1.png', width='300px', align='center', alt='sample alt')
    doctree += image

    doc = _walk_doctree(doctree, fake_builder)
    assert doc.inline_shapes, 'no inline shape created'
    shape = doc.inline_shapes[0]
    expected_width = 300 * 9525
    assert shape.width == expected_width, (shape.width, expected_width)

    # Alignment should have been applied to the picture's paragraph.
    picture_paragraphs = [p for p in doc.paragraphs if p.text == '']
    assert any(
        p.alignment is not None and p.alignment.value == 1  # CENTER
        for p in picture_paragraphs
    ), [(p.text, p.alignment) for p in picture_paragraphs]

    # Alt text propagates to the wp:docPr element.
    doc_pr = shape._inline.docPr  # noqa: SLF001
    assert doc_pr.get('descr') == 'sample alt'
    assert doc_pr.get('title') == 'sample alt'


def test_definition_list_renders_as_two_column_table(
    fake_builder: SimpleNamespace,
) -> None:
    """A three-item definition list produces a 2-col table with 3 rows; each
    term appears bold in col 0 and its definition text in col 1."""
    doctree = _blank_document()
    dl = nodes.definition_list()
    for term_text, def_text in [
        ('alpha', 'first definition'),
        ('beta', 'second definition'),
        ('gamma', 'third definition'),
    ]:
        item = nodes.definition_list_item()
        item += nodes.term('', term_text)
        definition = nodes.definition()
        definition += nodes.paragraph('', def_text)
        item += definition
        dl += item
    doctree += dl

    doc = _walk_doctree(doctree, fake_builder)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert len(table.columns) == 2
    expected = [
        ('alpha', 'first definition'),
        ('beta', 'second definition'),
        ('gamma', 'third definition'),
    ]
    for row, (want_term, want_def) in zip(table.rows, expected, strict=True):
        term_cell = row.cells[0]
        def_cell = row.cells[1]
        assert term_cell.paragraphs[0].text == want_term
        # Term is bold.
        assert any(r.bold for r in term_cell.paragraphs[0].runs), (
            want_term, [(r.text, r.bold) for r in term_cell.paragraphs[0].runs]
        )
        assert want_def in ' '.join(p.text for p in def_cell.paragraphs)


def test_definition_list_multi_definition_adds_paragraphs(
    fake_builder: SimpleNamespace,
) -> None:
    """MyST-style multiple ``<definition>`` children for one term land as
    consecutive paragraphs in column 1."""
    doctree = _blank_document()
    dl = nodes.definition_list()
    item = nodes.definition_list_item()
    item += nodes.term('', 'ambiguous term')
    for def_text in ('sense one', 'sense two', 'sense three'):
        d = nodes.definition()
        d += nodes.paragraph('', def_text)
        item += d
    dl += item
    doctree += dl

    doc = _walk_doctree(doctree, fake_builder)
    cell = doc.tables[0].rows[0].cells[1]
    paragraphs = [p.text for p in cell.paragraphs if p.text.strip()]
    assert paragraphs == ['sense one', 'sense two', 'sense three'], paragraphs


def test_option_group_signature_helper() -> None:
    """The flattener must combine option + option_argument into a single string
    per option, and join options in a group with ', '."""
    from docxsphinx.writer import _option_group_signature

    group = nodes.option_group()
    # -h option (no argument)
    opt_h = nodes.option()
    opt_h += nodes.option_string('', '-h')
    # --help option
    opt_help = nodes.option()
    opt_help += nodes.option_string('', '--help')
    group += opt_h
    group += opt_help
    assert _option_group_signature(group) == '-h, --help'

    # -o FILE, --out=FILE
    group2 = nodes.option_group()
    opt_short = nodes.option()
    opt_short += nodes.option_string('', '-o')
    opt_short += nodes.option_argument('', 'FILE', delimiter=' ')
    opt_long = nodes.option()
    opt_long += nodes.option_string('', '--out')
    opt_long += nodes.option_argument('', 'FILE', delimiter='=')
    group2 += opt_short
    group2 += opt_long
    assert _option_group_signature(group2) == '-o FILE, --out=FILE'


def test_option_list_renders_as_two_column_table(
    fake_builder: SimpleNamespace,
) -> None:
    """A three-option list produces a 2-col table with 3 rows."""
    doctree = _blank_document()
    ol = nodes.option_list()

    for sig_parts, desc_text in [
        ([('-h',)], 'Print help.'),
        ([('-v',), ('--verbose',)], 'Be verbose.'),
        ([('-o', 'FILE', ' '), ('--out', 'FILE', '=')], 'Write to FILE.'),
    ]:
        item = nodes.option_list_item()
        group = nodes.option_group()
        for parts in sig_parts:
            option = nodes.option()
            option += nodes.option_string('', parts[0])
            if len(parts) > 1:
                option += nodes.option_argument('', parts[1], delimiter=parts[2])
            group += option
        item += group
        desc = nodes.description()
        desc += nodes.paragraph('', desc_text)
        item += desc
        ol += item
    doctree += ol

    doc = _walk_doctree(doctree, fake_builder)
    assert len(doc.tables) == 1
    rows = doc.tables[0].rows
    assert len(rows) == 3
    assert rows[0].cells[0].paragraphs[0].text == '-h'
    assert rows[1].cells[0].paragraphs[0].text == '-v, --verbose'
    assert rows[2].cells[0].paragraphs[0].text == '-o FILE, --out=FILE'
    assert rows[2].cells[1].paragraphs[0].text == 'Write to FILE.'


def test_rowspan_two_rows_merges_cells(fake_builder: SimpleNamespace) -> None:
    """A 2x2 table where the first-column cell has ``morerows=1`` must emit
    a ``vMerge`` restart on the top cell and a continuation on the bottom."""
    doctree = _blank_document()
    table = nodes.table()
    tgroup = nodes.tgroup(cols=2)
    tgroup += nodes.colspec(colwidth=10)
    tgroup += nodes.colspec(colwidth=10)
    tbody = nodes.tbody()

    # Row 0: entry(morerows=1) | entry
    row0 = nodes.row()
    e0 = nodes.entry(morerows=1)
    e0 += nodes.paragraph('', 'spanning')
    e1 = nodes.entry()
    e1 += nodes.paragraph('', 'top right')
    row0 += e0
    row0 += e1

    # Row 1: just one entry — column 0 is occupied by the rowspan.
    row1 = nodes.row()
    e2 = nodes.entry()
    e2 += nodes.paragraph('', 'bottom right')
    row1 += e2

    tbody += row0
    tbody += row1
    tgroup += tbody
    table += tgroup
    doctree += table

    doc = _walk_doctree(doctree, fake_builder)
    assert len(doc.tables) == 1
    tbl = doc.tables[0]

    from docx.oxml.ns import qn

    def vmerge_layout(table):
        out = []
        for tr in table._tbl.findall(qn('w:tr')):
            row_out = []
            for tc in tr.findall(qn('w:tc')):
                tcPr = tc.find(qn('w:tcPr'))
                v = tcPr.find(qn('w:vMerge')) if tcPr is not None else None
                if v is None:
                    row_out.append(None)
                else:
                    row_out.append(v.get(qn('w:val')) or 'continue')
            out.append(row_out)
        return out

    layout = vmerge_layout(tbl)
    # Row 0: [spanning (restart), top right (no vmerge)]
    # Row 1: [continuation (continue), bottom right (no vmerge)]
    assert layout == [['restart', None], ['continue', None]], layout

    # The continuation cell has no text of its own (merge discards its content);
    # the right-side cell on row 1 must still carry its content.
    # Iterate raw tc to avoid the python-docx cell-collapse.
    trs = tbl._tbl.findall(qn('w:tr'))
    row1_tcs = trs[1].findall(qn('w:tc'))
    row1_right_text = ''.join(
        t.text or ''
        for t in row1_tcs[1].iter(qn('w:t'))
    )
    assert row1_right_text == 'bottom right', row1_right_text


def test_external_hyperlink_creates_hyperlink_element_and_relationship(
    fake_builder: SimpleNamespace,
) -> None:
    """External-URL refs produce a ``w:hyperlink`` element whose ``r:id``
    resolves via the document's relationships to the URL, with the link
    text preserved."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    para = nodes.paragraph()
    para += nodes.Text('Visit ')
    ref = nodes.reference(refuri='https://www.sphinx-doc.org/')
    ref += nodes.Text('Sphinx')
    para += ref
    para += nodes.Text(' today.')
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    paragraph = doc.paragraphs[0]

    hyperlinks = paragraph._p.findall(qn('w:hyperlink'))
    assert len(hyperlinks) == 1, paragraph._p.xml
    link = hyperlinks[0]
    r_id = link.get(qn('r:id'))
    assert r_id, link.attrib
    # Visible text is the anchor text from the reference node.
    text_elems = link.findall('.//' + qn('w:t'))
    assert ''.join(t.text for t in text_elems) == 'Sphinx'
    # The relationship resolves to the URL.
    rels = paragraph.part.rels
    assert rels[r_id].target_ref == 'https://www.sphinx-doc.org/'
    assert rels[r_id].is_external is True


def test_relative_url_reference_falls_back_to_plain_text(
    fake_builder: SimpleNamespace,
) -> None:
    """Relative ``refuri`` without a scheme is rendered as plain text —
    bookmark anchoring is a later phase."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    para = nodes.paragraph()
    ref = nodes.reference(refuri='./other.html')
    ref += nodes.Text('relative link')
    para += ref
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    paragraph = doc.paragraphs[0]

    assert not paragraph._p.findall(qn('w:hyperlink')), paragraph._p.xml
    assert 'relative link' in paragraph.text


def test_internal_refid_reference_becomes_internal_hyperlink(
    fake_builder: SimpleNamespace,
) -> None:
    """A reference carrying only ``refid`` renders as a ``w:hyperlink`` with
    ``w:anchor=<refid>`` (not an external ``r:id``)."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    para = nodes.paragraph()
    ref = nodes.reference(refid='my-section')
    ref += nodes.Text('my section')
    para += ref
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    paragraph = doc.paragraphs[0]

    hyperlinks = paragraph._p.findall(qn('w:hyperlink'))
    assert len(hyperlinks) == 1
    link = hyperlinks[0]
    assert link.get(qn('w:anchor')) == 'my-section'
    assert link.get(qn('r:id')) is None
    text_elems = link.findall('.//' + qn('w:t'))
    assert ''.join(t.text for t in text_elems) == 'my section'


def test_section_emits_bookmark_for_its_ids(
    fake_builder: SimpleNamespace,
) -> None:
    """A ``section`` node with ``ids=['my-id']`` emits a ``w:bookmarkStart/End``
    pair with ``name='my-id'`` in the heading paragraph, enabling internal
    references to anchor at the section."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    section = nodes.section(ids=['my-id'])
    section += nodes.title('', 'My heading')
    section += nodes.paragraph('', 'body content')
    doctree += section

    doc = _walk_doctree(doctree, fake_builder)
    # Find the heading paragraph.
    heading = doc.paragraphs[0]
    bookmarks = heading._p.findall(qn('w:bookmarkStart'))
    assert len(bookmarks) == 1
    assert bookmarks[0].get(qn('w:name')) == 'my-id'
    # And a matching bookmarkEnd with the same id.
    ends = heading._p.findall(qn('w:bookmarkEnd'))
    assert len(ends) == 1
    assert ends[0].get(qn('w:id')) == bookmarks[0].get(qn('w:id'))


def test_intra_doc_fragment_refuri_becomes_internal_hyperlink(
    fake_builder: SimpleNamespace,
) -> None:
    """``refuri='#target'`` should be parsed as an internal anchor."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    para = nodes.paragraph()
    ref = nodes.reference(refuri='#target-section')
    ref += nodes.Text('jump')
    para += ref
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    paragraph = doc.paragraphs[0]
    hyperlinks = paragraph._p.findall(qn('w:hyperlink'))
    assert len(hyperlinks) == 1
    assert hyperlinks[0].get(qn('w:anchor')) == 'target-section'


def test_python_code_block_produces_multiple_colored_runs(
    fake_builder: SimpleNamespace,
) -> None:
    """A Python literal_block emits one paragraph with several colour-differentiated
    runs, including at least two distinct RGB colours (keyword, string, number,
    comment, …). A single plain-text run is the pre-Phase-2.3 failure mode."""
    doctree = _blank_document()
    code = (
        'def greet(name):\n'
        '    """Say hi."""\n'
        "    return f'hello, {name}!'\n"
    )
    lb = nodes.literal_block(code, code, language='python')
    doctree += lb

    doc = _walk_doctree(doctree, fake_builder)
    assert len(doc.paragraphs) == 1
    runs = doc.paragraphs[0].runs
    assert len(runs) > 1, [r.text for r in runs]  # fails loudly on plain-text regression

    colors = {
        str(r.font.color.rgb)
        for r in runs
        if r.font.color is not None and r.font.color.rgb is not None
    }
    assert len(colors) >= 2, colors

    # Every run with actual (non-whitespace) text should carry the code font.
    # Break runs emitted for newlines return ``'\n'`` from ``run.text`` — skip
    # those explicitly.
    text_runs = [r for r in runs if r.text and r.text.strip()]
    assert text_runs, runs
    assert all(r.font.name == 'Consolas' for r in text_runs), [
        (r.text, r.font.name) for r in text_runs
    ]


def test_unknown_language_falls_back_to_plain_monospace(
    fake_builder: SimpleNamespace,
) -> None:
    """A literal_block with no language (or an unknown one) renders as plain
    monospace — runs still get the code font but no colour is applied."""
    doctree = _blank_document()
    lb = nodes.literal_block('plain code line 1\nplain code line 2\n', '')
    doctree += lb

    doc = _walk_doctree(doctree, fake_builder)
    runs = doc.paragraphs[0].runs
    assert any(r.text for r in runs)
    coloured = [
        r for r in runs
        if r.font.color is not None and r.font.color.rgb is not None
    ]
    assert coloured == [], [(r.text, r.font.color.rgb) for r in coloured]


def test_code_block_preserves_line_breaks(
    fake_builder: SimpleNamespace,
) -> None:
    """Newlines in the source must turn into ``<w:br/>`` elements inside the
    single code paragraph."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    code = 'line one\nline two\nline three\n'
    lb = nodes.literal_block(code, code, language='text')
    doctree += lb

    doc = _walk_doctree(doctree, fake_builder)
    breaks = doc.paragraphs[0]._p.findall('.//' + qn('w:br'))
    # Three newlines in the source → three line breaks.
    assert len(breaks) == 3, f'expected 3 w:br, got {len(breaks)}'


def _build_figure_doctree_for_caption(fid: str | None, caption_text: str) -> nodes.document:
    doctree = _blank_document()
    fig = nodes.figure(ids=[fid] if fid else [])
    fig += nodes.caption('', caption_text)
    doctree += fig
    return doctree


def test_figure_caption_emits_caption_paragraph_with_seq_field_and_bookmark(
    fake_builder: SimpleNamespace,
) -> None:
    """A single ``figure`` with a caption produces one paragraph with the
    ``Caption`` style (when available), a ``SEQ Figure`` field for Word
    auto-numbering, a bookmark pointing at the figure's id, and the caption
    text tacked onto the end."""
    from docx.oxml.ns import qn

    doctree = _build_figure_doctree_for_caption('my-fig', 'A pair of cats.')
    doc = _walk_doctree(doctree, fake_builder)

    caption_paragraphs = [p for p in doc.paragraphs if p.text and 'Figure' in p.text]
    assert len(caption_paragraphs) == 1, [(p.text, p.style.name) for p in doc.paragraphs]
    p = caption_paragraphs[0]

    # SEQ field present, naming the Figure sequence.
    flds = p._p.findall(qn('w:fldSimple'))
    assert len(flds) == 1
    assert 'SEQ Figure' in flds[0].get(qn('w:instr'))

    # A bookmark with the figure's id is emitted in the caption paragraph.
    bms = p._p.findall(qn('w:bookmarkStart'))
    assert any(b.get(qn('w:name')) == 'my-fig' for b in bms), [
        b.get(qn('w:name')) for b in bms
    ]

    # Caption text is appended after the prefix.
    assert 'A pair of cats.' in p.text


def test_sequential_figures_get_sequential_numbers(
    fake_builder: SimpleNamespace,
) -> None:
    """Each figure bumps the translator's ``_figure_counter``; the number
    passed to the SEQ field placeholder matches the 1-based figure index."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    for fid, text in [('fig-1', 'First figure'), ('fig-2', 'Second figure'), ('fig-3', 'Third')]:
        fig = nodes.figure(ids=[fid])
        fig += nodes.caption('', text)
        doctree += fig

    doc = _walk_doctree(doctree, fake_builder)
    # Each caption paragraph has a SEQ field whose placeholder text is the
    # expected figure number.
    seq_placeholders = []
    for p in doc.paragraphs:
        for fld in p._p.findall(qn('w:fldSimple')):
            texts = [t.text for t in fld.findall('.//' + qn('w:t'))]
            seq_placeholders.append(''.join(texts))
    assert seq_placeholders == ['1', '2', '3'], seq_placeholders


def test_footnotes_part_is_created_on_first_footnote(
    fake_builder: SimpleNamespace, tmp_path,
) -> None:
    """Writing a doctree containing a footnote reference + body must add a
    ``word/footnotes.xml`` part to the saved package, with the right number
    of ``<w:footnote>`` entries (two separators + one per user footnote)."""
    import zipfile

    doctree = _blank_document()
    # <paragraph>[1]</paragraph> where the ref targets footnote-1
    para = nodes.paragraph()
    para += nodes.Text('See ')
    ref = nodes.footnote_reference(refid='footnote-1')
    ref += nodes.Text('1')
    para += ref
    para += nodes.Text(' for details.')
    doctree += para

    footnote = nodes.footnote(ids=['footnote-1'])
    footnote += nodes.label('', '1')
    footnote += nodes.paragraph('', 'This is the footnote body.')
    doctree += footnote

    doc = _walk_doctree(doctree, fake_builder)
    out_path = tmp_path / 'out.docx'
    doc.save(str(out_path))

    with zipfile.ZipFile(out_path) as zf:
        assert 'word/footnotes.xml' in zf.namelist()
        xml = zf.read('word/footnotes.xml').decode()
        # Two separator footnotes + one user footnote = 3 <w:footnote> elements.
        assert xml.count('<w:footnote ') == 3, xml
        # The body text shows up in the user footnote.
        assert 'This is the footnote body.' in xml


def test_footnote_reference_and_body_share_ooxml_id(
    fake_builder: SimpleNamespace,
) -> None:
    """The OOXML id emitted on the ``<w:footnoteReference>`` inline element
    must match the id used when the ``<w:footnote>`` body entry is written
    — otherwise Word shows no footnote content at the reference."""
    from docx.oxml.ns import qn

    doctree = _blank_document()

    # ref 1 → footnote 'fn-a' ; ref 2 → footnote 'fn-b'
    para = nodes.paragraph()
    r1 = nodes.footnote_reference(refid='fn-a')
    r1 += nodes.Text('1')
    para += r1
    para += nodes.Text(' then ')
    r2 = nodes.footnote_reference(refid='fn-b')
    r2 += nodes.Text('2')
    para += r2
    doctree += para

    fn_a = nodes.footnote(ids=['fn-a'])
    fn_a += nodes.paragraph('', 'Body A')
    doctree += fn_a

    fn_b = nodes.footnote(ids=['fn-b'])
    fn_b += nodes.paragraph('', 'Body B')
    doctree += fn_b

    doc = _walk_doctree(doctree, fake_builder)

    # Collect the ids of inline footnote references in order.
    ref_ids = []
    for p in doc.paragraphs:
        for ref in p._p.findall('.//' + qn('w:footnoteReference')):
            ref_ids.append(ref.get(qn('w:id')))
    # Should be ['1', '2'] — the first encountered gets id 1, the second 2.
    assert ref_ids == ['1', '2'], ref_ids


def test_citation_reference_uses_same_mechanism(
    fake_builder: SimpleNamespace,
) -> None:
    """Citations and citation references share the footnote pipeline."""
    doctree = _blank_document()
    para = nodes.paragraph()
    ref = nodes.citation_reference(refid='cit-1')
    ref += nodes.Text('Smith')
    para += ref
    doctree += para

    cit = nodes.citation(ids=['cit-1'])
    cit += nodes.label('', 'Smith')
    cit += nodes.paragraph('', 'Smith 2024')
    doctree += cit

    doc = _walk_doctree(doctree, fake_builder)

    # Save and inspect the resulting zip — citation body should live in
    # footnotes.xml just like footnotes do.
    import io
    import zipfile
    buf = io.BytesIO()
    doc.save(buf)
    with zipfile.ZipFile(buf) as zf:
        fn_xml = zf.read('word/footnotes.xml').decode()
    assert 'Smith 2024' in fn_xml


def test_strikethrough_text_sets_run_strike(fake_builder: SimpleNamespace) -> None:
    """MyST's strikethrough emits raw-HTML ``<s>`` / ``</s>`` sibling
    markers. The writer toggles a strike flag on those, so the text
    between them renders with ``run.font.strike = True``."""
    doctree = _blank_document()
    para = nodes.paragraph()
    para += nodes.Text('before ')
    para += nodes.raw('', '<s>', format='html')
    para += nodes.Text('struck')
    para += nodes.raw('', '</s>', format='html')
    para += nodes.Text(' after')
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    runs = doc.paragraphs[0].runs
    struck = [r for r in runs if r.font.strike]
    assert any('struck' in r.text for r in struck), [
        (r.text, r.font.strike) for r in runs
    ]
    # Text outside the markers is not struck.
    unstruck = [r for r in runs if not r.font.strike]
    assert any('before' in r.text for r in unstruck)
    assert any('after' in r.text for r in unstruck)


def test_hyperlink_preserves_inline_bold(fake_builder: SimpleNamespace) -> None:
    """``[**bold link**](url)`` — the inner ``<strong>`` is preserved as a
    single bold run *inside* the ``<w:hyperlink>`` element."""
    from docx.oxml.ns import qn

    doctree = _blank_document()
    para = nodes.paragraph()
    ref = nodes.reference(refuri='https://example.com/')
    strong = nodes.strong()
    strong += nodes.Text('bold link')
    ref += strong
    para += ref
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    hyperlinks = doc.paragraphs[0]._p.findall(qn('w:hyperlink'))
    assert len(hyperlinks) == 1
    runs = hyperlinks[0].findall(qn('w:r'))
    assert len(runs) == 1, len(runs)
    rPr = runs[0].find(qn('w:rPr'))
    assert rPr is not None
    assert rPr.find(qn('w:b')) is not None, 'hyperlink run should be bold'
    # Confirm the link anchor text is correct.
    t = runs[0].find('.//' + qn('w:t'))
    assert t.text == 'bold link'


def test_footnote_body_preserves_inline_bold(
    fake_builder: SimpleNamespace, tmp_path,
) -> None:
    """A footnote body containing ``**bold**`` emits a bold run inside the
    ``<w:footnote>`` entry — not just flattened plain text."""
    import zipfile


    doctree = _blank_document()
    # In-body reference
    para = nodes.paragraph()
    ref = nodes.footnote_reference(refid='fn-1')
    ref += nodes.Text('1')
    para += ref
    doctree += para

    # Footnote body: "Plain **bold** more."
    fn = nodes.footnote(ids=['fn-1'])
    fn += nodes.label('', '1')
    body_para = nodes.paragraph()
    body_para += nodes.Text('Plain ')
    strong = nodes.strong()
    strong += nodes.Text('bold')
    body_para += strong
    body_para += nodes.Text(' more.')
    fn += body_para
    doctree += fn

    doc = _walk_doctree(doctree, fake_builder)
    out = tmp_path / 'fn.docx'
    doc.save(str(out))
    with zipfile.ZipFile(out) as zf:
        xml = zf.read('word/footnotes.xml').decode()
    # There must be a bold run inside the user footnote body.
    assert '<w:b/>' in xml or '<w:b />' in xml, xml[:600]
    assert 'Plain' in xml and 'bold' in xml and 'more' in xml


def test_inline_math_renders_as_omml_in_current_paragraph(
    fake_builder: SimpleNamespace,
) -> None:
    """``:math:`E = mc^2``` → a ``<m:oMath>`` element appended to the
    containing paragraph, not a new paragraph of its own."""

    doctree = _blank_document()
    para = nodes.paragraph()
    para += nodes.Text('Einstein: ')
    math = nodes.math()
    math += nodes.Text('E = mc^2')
    para += math
    para += nodes.Text('.')
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    # One paragraph — the math goes inline, not into its own block.
    assert len(doc.paragraphs) == 1, [p.text for p in doc.paragraphs]
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    omath = doc.paragraphs[0]._p.findall(f'.//{{{M}}}oMath')
    assert len(omath) == 1, 'expected one <m:oMath> inside the paragraph'
    # The trailing '.' text run is still present too.
    assert 'Einstein' in doc.paragraphs[0].text
    assert '.' in doc.paragraphs[0].text


def test_display_math_block_renders_as_oMathPara_in_new_paragraph(
    fake_builder: SimpleNamespace,
) -> None:
    """``.. math::`` → a fresh paragraph containing ``<m:oMathPara>``."""
    doctree = _blank_document()
    mb = nodes.math_block()
    mb += nodes.Text(r'\sum_{i=1}^n i = \frac{n(n+1)}{2}')
    doctree += mb

    doc = _walk_doctree(doctree, fake_builder)
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    paras_with_math = [
        p for p in doc.paragraphs
        if p._p.findall(f'.//{{{M}}}oMathPara')
    ]
    assert len(paras_with_math) == 1, (
        'expected exactly one paragraph with <m:oMathPara>'
    )


def test_math_falls_back_to_monospace_when_pandoc_missing(
    fake_builder: SimpleNamespace,
) -> None:
    """If pandoc isn't on PATH, the writer emits the raw LaTeX as a
    Consolas-styled run instead of dropping or crashing."""
    from unittest.mock import patch

    doctree = _blank_document()
    para = nodes.paragraph()
    math = nodes.math()
    math += nodes.Text('alpha + beta')
    para += math
    doctree += para

    with patch('docxsphinx._docx_helpers.shutil.which', return_value=None):
        # Also bypass the lru_cache so the mock actually fires.
        from docxsphinx._docx_helpers import _latex_to_omml_bytes
        _latex_to_omml_bytes.cache_clear()
        doc = _walk_doctree(doctree, fake_builder)

    # LaTeX text made it into the document as plain text, monospace font.
    runs = [r for p in doc.paragraphs for r in p.runs]
    latex_runs = [r for r in runs if r.text == 'alpha + beta']
    assert latex_runs, [r.text for r in runs]
    assert latex_runs[0].font.name == 'Consolas', latex_runs[0].font.name


def test_empty_comment_does_not_raise_index_error(
    fake_builder: SimpleNamespace,
) -> None:
    """Regression for GitHub issues #33 and #55 — ``visit_comment`` used to
    index ``node[0]`` unconditionally, raising ``IndexError`` on an empty
    ``.. \\n`` comment (e.g. the one Sphinx's ``sphinx-quickstart`` inserts
    at the top of ``conf.py`` during env serialisation). Walking a doctree
    with a child-less ``comment`` node must now be a silent no-op."""
    doctree = _blank_document()
    doctree += nodes.comment()  # empty — no children
    doc = _walk_doctree(doctree, fake_builder)
    # Nothing should be emitted for the comment itself.
    assert doc.paragraphs == [] or all(not p.text for p in doc.paragraphs)


def test_comment_docxtablestyle_override_still_works(
    fake_builder: SimpleNamespace,
) -> None:
    """The ``DocxTableStyle <name>`` comment-directive hack must keep
    working after the IndexError fix."""
    doctree = _blank_document()
    comment = nodes.comment('', 'DocxTableStyle Light Grid Accent 1')
    doctree += comment
    # Build a minimal table afterwards to observe the table_style attribute.
    # (Direct assertion on translator state is cleaner — run the walker and
    # inspect.)
    from docx import Document

    from docxsphinx.writer import DocxTranslator

    container = Document()
    tr = DocxTranslator(doctree, fake_builder, container)
    doctree.walkabout(tr)
    assert tr.current_state.table_style == 'Light Grid Accent 1'


def test_generic_admonition_uses_custom_title(
    fake_builder: SimpleNamespace,
) -> None:
    """``.. admonition:: My Heading`` should put the user's heading in the
    title run instead of a default label."""
    doctree = _blank_document()
    adm = nodes.admonition()
    adm += nodes.title('', 'My custom heading')
    adm += nodes.paragraph('', 'Body text.')
    doctree += adm

    doc = _walk_doctree(doctree, fake_builder)
    cell = doc.tables[0].rows[0].cells[0]
    assert cell.paragraphs[0].text == 'My custom heading', cell.paragraphs[0].text
    # Confirm the title is also not re-rendered as a heading further down
    # in the cell — the node.remove() in visit_admonition prevents that.
    following = [p.text for p in cell.paragraphs[1:] if p.text.strip()]
    assert 'My custom heading' not in following, following


def test_field_list_renders_as_two_column_table(
    fake_builder: SimpleNamespace,
) -> None:
    """``:param x: …`` field lists (common in docstrings via autodoc) must
    emit a 2-column table with bold field names in the left column and the
    field body in the right."""
    doctree = _blank_document()
    fl = nodes.field_list()
    for name, body in (
        ('param name', 'The person to greet.'),
        ('returns', 'A greeting string.'),
    ):
        field = nodes.field()
        field += nodes.field_name(text=name)
        field_body = nodes.field_body()
        field_body += nodes.paragraph(text=body)
        field += field_body
        fl += field
    doctree += fl

    doc = _walk_doctree(doctree, fake_builder)
    assert len(doc.tables) == 1, [t for t in doc.tables]
    table = doc.tables[0]
    assert len(table.rows) == 2, [r.cells[0].text for r in table.rows]
    assert len(table.columns) == 2
    # Left column: field names, bold
    names = [row.cells[0].text for row in table.rows]
    assert 'param name' in names[0] or 'name' in names[0], names
    assert 'returns' in names[1].lower(), names
    for row in table.rows:
        runs = row.cells[0].paragraphs[0].runs
        assert runs and any(r.bold for r in runs), (
            'field name should be bold',
            [(r.text, r.bold) for r in runs],
        )
    # Right column: field bodies
    bodies = [row.cells[1].text for row in table.rows]
    assert 'person to greet' in bodies[0].lower(), bodies
    assert 'greeting string' in bodies[1].lower(), bodies


def test_desc_signature_renders_function_with_parameters(
    fake_builder: SimpleNamespace,
) -> None:
    """A synthetic Sphinx ``desc`` / ``desc_signature`` / ``desc_name`` /
    ``desc_parameterlist`` tree (the shape autodoc builds) should emit a
    paragraph containing ``name(param1, param2)`` with the name in bold."""
    from sphinx import addnodes

    doctree = _blank_document()
    desc = addnodes.desc()
    desc['domain'] = 'py'
    desc['objtype'] = 'function'
    sig = addnodes.desc_signature('', '')
    sig += addnodes.desc_name('', 'greet')
    plist = addnodes.desc_parameterlist()
    plist += addnodes.desc_parameter('', 'name')
    plist += addnodes.desc_parameter('', 'greeting')
    sig += plist
    sig += addnodes.desc_returns('', 'str')
    desc += sig
    desc += addnodes.desc_content()
    doctree += desc

    doc = _walk_doctree(doctree, fake_builder)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    sig_line = next((t for t in texts if 'greet' in t), None)
    assert sig_line is not None, texts
    assert 'greet(name, greeting)' in sig_line, sig_line
    assert '→' in sig_line and 'str' in sig_line, sig_line

    # The function name run should be bold.
    bold_runs = [
        r.text for p in doc.paragraphs for r in p.runs
        if r.bold and 'greet' in r.text
    ]
    assert bold_runs, [(r.text, r.bold) for p in doc.paragraphs for r in p.runs]


def test_desc_optional_emits_brackets(fake_builder: SimpleNamespace) -> None:
    """``desc_optional`` (used for ``f(a[, b[, c]])`` style signatures)
    must wrap its contents in literal square brackets."""
    from sphinx import addnodes

    doctree = _blank_document()
    desc = addnodes.desc()
    desc['domain'] = 'py'
    desc['objtype'] = 'function'
    sig = addnodes.desc_signature('', '')
    sig += addnodes.desc_name('', 'func')
    plist = addnodes.desc_parameterlist()
    plist += addnodes.desc_parameter('', 'a')
    optional = addnodes.desc_optional()
    optional += addnodes.desc_parameter('', 'b')
    plist += optional
    sig += plist
    desc += sig
    doctree += desc

    doc = _walk_doctree(doctree, fake_builder)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    sig_line = next((t for t in texts if 'func' in t), None)
    assert sig_line is not None, texts
    # Optional parameters are wrapped in literal '[' / ']' — exact inner
    # punctuation varies with parameter-list join logic (may render as
    # '[, b]' or '[b]' depending on whether the optional is first).
    assert '[' in sig_line and ']' in sig_line, sig_line
    assert 'func(a' in sig_line and 'b' in sig_line, sig_line


def test_desc_addname_module_prefix_hidden(fake_builder: SimpleNamespace) -> None:
    """``desc_addname`` carries the module prefix (``mymodule.``); we
    deliberately suppress it to keep signatures compact in Word."""
    from sphinx import addnodes

    doctree = _blank_document()
    desc = addnodes.desc()
    desc['domain'] = 'py'
    desc['objtype'] = 'class'
    sig = addnodes.desc_signature('', '')
    sig += addnodes.desc_addname('', 'mymodule.')
    sig += addnodes.desc_name('', 'Calculator')
    sig += addnodes.desc_parameterlist()
    desc += sig
    doctree += desc

    doc = _walk_doctree(doctree, fake_builder)
    all_text = ' '.join(p.text for p in doc.paragraphs)
    assert 'Calculator' in all_text, all_text
    assert 'mymodule.' not in all_text, all_text


def _make_template_builder(
    srcdir: Path, docx_template: str, templates_path: list[str] | None = None,
) -> SimpleNamespace:
    """Minimal ``Builder`` stand-in wired up for ``DocxWriter`` to resolve
    a ``docx_template`` against ``srcdir`` + optional ``templates_path``."""
    cfg: dict = {'docx_template': docx_template, 'docx_debug_log': None}
    if templates_path is not None:
        cfg['templates_path'] = templates_path
    return SimpleNamespace(env=SimpleNamespace(srcdir=str(srcdir)), config=cfg)


def test_template_path_resolves_against_srcdir(tmp_path: Path) -> None:
    """Back-compat: ``docx_template = 'template.docx'`` in ``<srcdir>/`` still
    resolves when ``templates_path`` is unset."""
    from docxsphinx.writer import DocxWriter

    template = tmp_path / 'template.docx'
    # Minimal valid docx — copy the python-docx default.
    Document().save(str(template))

    builder = _make_template_builder(tmp_path, 'template.docx')
    writer = DocxWriter(builder)
    assert writer.template_path == template


def test_template_path_honours_templates_path(tmp_path: Path) -> None:
    """``templates_path`` entries are searched before ``srcdir``, fixing #22."""
    from docxsphinx.writer import DocxWriter

    templates_dir = tmp_path / '_templates'
    templates_dir.mkdir()
    template = templates_dir / 'template.docx'
    Document().save(str(template))

    builder = _make_template_builder(
        tmp_path, 'template.docx', templates_path=['_templates'],
    )
    writer = DocxWriter(builder)
    assert writer.template_path == template


def test_template_path_absolute_path_untouched(tmp_path: Path) -> None:
    """Absolute ``docx_template`` is used as-is regardless of
    ``templates_path``."""
    from docxsphinx.writer import DocxWriter

    template = tmp_path / 'abs_template.docx'
    Document().save(str(template))

    builder = _make_template_builder(
        tmp_path, str(template), templates_path=['_templates'],
    )
    writer = DocxWriter(builder)
    assert writer.template_path == template


def test_numref_reference_renders_as_hyperlink(
    fake_builder: SimpleNamespace,
) -> None:
    """`:numref:` resolves to a ``docutils.nodes.number_reference`` wrapping
    an inline with formatted text like "Fig. 1". Before the fix for #36,
    no ``visit_number_reference`` existed so ``unknown_visit`` silently
    skipped the whole node — references like "numbered reference Fig. 1"
    rendered as "numbered reference ". Now it routes through the same
    internal-hyperlink emitter as regular ``:ref:``."""
    from docx.oxml.ns import qn
    from sphinx import addnodes

    doctree = _blank_document()
    para = nodes.paragraph()
    para += nodes.Text('see ')
    numref = addnodes.number_reference(
        '', '', refid='first-figure', title='Fig. %s',
    )
    numref['internal'] = True
    inner = nodes.inline(classes=['std', 'std-numref'])
    inner += nodes.Text('Fig. 1')
    numref += inner
    para += numref
    doctree += para

    doc = _walk_doctree(doctree, fake_builder)
    # The paragraph should contain a hyperlink to the figure anchor with
    # the resolved "Fig. 1" text.
    hyperlinks = doc.paragraphs[0]._element.findall('.//' + qn('w:hyperlink'))
    assert len(hyperlinks) == 1, [p.text for p in doc.paragraphs]
    hl = hyperlinks[0]
    assert hl.get(qn('w:anchor')) == 'first-figure'
    text = ''.join(t.text or '' for t in hl.findall('.//' + qn('w:t')))
    assert text == 'Fig. 1', text


def test_template_path_templates_path_takes_precedence(tmp_path: Path) -> None:
    """When the same filename exists in both a ``templates_path`` entry
    and ``srcdir``, the ``templates_path`` copy wins (users who set
    ``templates_path`` expect Sphinx's standard template lookup)."""
    from docxsphinx.writer import DocxWriter

    templates_dir = tmp_path / '_templates'
    templates_dir.mkdir()
    in_templates = templates_dir / 'template.docx'
    in_srcdir = tmp_path / 'template.docx'
    Document().save(str(in_templates))
    Document().save(str(in_srcdir))

    builder = _make_template_builder(
        tmp_path, 'template.docx', templates_path=['_templates'],
    )
    writer = DocxWriter(builder)
    assert writer.template_path == in_templates
