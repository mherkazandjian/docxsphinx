"""Custom docutils writer for OpenXML (.docx).

Originally based on sphinxcontrib-docxbuilder (copyright 2010 shimizukawa,
BSD licence).
"""
from __future__ import annotations

import logging
import re
import sys
from pathlib import Path

from docutils import nodes, writers
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor
from pygments import lex
from pygments.lexers import TextLexer, get_lexer_by_name
from pygments.token import Token
from pygments.util import ClassNotFound

from docxsphinx._compat import _Cell
from docxsphinx._docx_helpers import (
    add_bookmark,
    add_footnote,
    add_footnote_reference,
    add_hyperlink,
    add_internal_hyperlink,
    add_seq_field,
    ensure_footnotes_part,
    latex_to_omml,
)

logger = logging.getLogger(__name__)


BALLOT_BOX = '☐'           # ☐ empty
BALLOT_BOX_WITH_X = '☒'    # ☒ checked


def _checkbox_prefix(list_item: nodes.Element) -> str:
    """Return ``☒`` or ``☐`` for a MyST ``task-list-item`` list_item node.

    MyST emits the checkbox as a raw-HTML ``<input>`` element inside the
    item's first paragraph; the presence of ``checked`` in that raw text
    determines the glyph. No trailing space is returned — MyST already
    leaves a leading space in the text node that follows the raw element,
    which provides natural separation.
    """
    checked = False
    for raw in list_item.findall(nodes.raw):
        if raw.get('format') != 'html':
            continue
        if 'checked' in raw.astext():
            checked = True
        break
    return BALLOT_BOX_WITH_X if checked else BALLOT_BOX


def dprint(_func: str | None = None, **kw: object) -> None:
    """Emit a DEBUG-level trace line for the current visitor method.

    Called with no kwargs (the common case) this is a no-op — enable debug
    tracing by calling with explicit ``key=value`` pairs from inside a
    ``visit_*``/``depart_*`` method, or by setting the ``docx_debug_log``
    config value and raising this logger to DEBUG in ``conf.py``.
    """
    if not kw:
        return
    frame = sys._getframe(1)
    func = _func or frame.f_code.co_name
    text = ', '.join(f'{k} = {v}' for k, v in kw.items())
    logger.debug('%s %s', func, text)


_LENGTH_RE = re.compile(r'^\s*([0-9]*\.?[0-9]+)\s*([a-zA-Z%]*)\s*$')

# Schemes that reliably identify an *external* URL for hyperlink rendering.
# Relative/internal refs (no scheme, or ``refid`` only) are left as plain
# text until bookmark support lands.
_EXTERNAL_URL_RE = re.compile(r'^(https?|ftp|mailto|file|data|tel):', re.IGNORECASE)

# 1 CSS pixel at the conventional 96 dpi = 9525 English Metric Units.
_EMU_PER_PX = 9525


def _length_to_emu(value: str | None):
    """Convert a CSS-ish length (``"300px"``, ``"5cm"``, ``"2in"``, ``"72pt"``)
    into a python-docx length (an :class:`Emu` subclass). Returns ``None`` for
    percentage inputs or unrecognised units — callers should fall back to
    python-docx's default sizing in that case.
    """
    if not value:
        return None
    match = _LENGTH_RE.match(value)
    if not match:
        return None
    num = float(match.group(1))
    unit = match.group(2).lower()
    if unit in ('', 'px'):
        return Emu(int(num * _EMU_PER_PX))
    if unit == 'cm':
        return Cm(num)
    if unit == 'mm':
        return Mm(num)
    if unit == 'in':
        return Inches(num)
    if unit == 'pt':
        return Pt(num)
    return None  # unknown unit ('%', 'em', 'ex', …) — defer to native sizing


_ALIGN_MAP = {
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'left':   WD_ALIGN_PARAGRAPH.LEFT,
    'right':  WD_ALIGN_PARAGRAPH.RIGHT,
}


CODE_FONT_NAME = 'Consolas'
"""Monospace face applied to every run emitted for a code block. Consolas ships with Windows by default; macOS/Linux Word substitutes with a similar monospace if the font is missing."""

# Pygments token → (hex color, bold, italic). Lookup walks up the token
# hierarchy (Token.Keyword.Constant → Token.Keyword → Token), so broad
# families only need one entry.
_TOKEN_STYLE: dict[object, tuple[str, bool, bool]] = {
    Token.Keyword:              ('00007F', True,  False),   # dark blue, bold
    Token.Name.Builtin:         ('008080', False, False),   # teal
    Token.Name.Function:        ('7F0055', False, False),   # purple
    Token.Name.Class:           ('7F0055', True,  False),
    Token.Name.Decorator:       ('7F0055', False, False),
    Token.Name.Namespace:       ('7F0055', False, False),
    Token.Name.Tag:             ('800080', False, False),
    Token.Name.Attribute:       ('7F7F00', False, False),
    Token.Name.Exception:       ('7F0055', True,  False),
    Token.Name.Variable:        ('000000', False, False),
    Token.String:               ('2A7F00', False, False),   # green
    Token.String.Doc:           ('2A7F00', False, True),
    Token.String.Escape:        ('CC7F00', False, False),
    Token.Number:               ('A52A2A', False, False),   # brown
    Token.Comment:              ('808080', False, True),    # gray italic
    Token.Comment.Preproc:      ('7F0055', False, False),
    Token.Operator:             ('000000', False, False),
    Token.Operator.Word:        ('00007F', True,  False),
    Token.Punctuation:          ('000000', False, False),
    Token.Generic.Heading:      ('000080', True,  False),
    Token.Generic.Subheading:   ('000080', False, False),
    Token.Generic.Deleted:      ('A52A2A', False, False),
    Token.Generic.Inserted:     ('00AA00', False, False),
    Token.Generic.Error:        ('FF0000', False, False),
}


def _token_style(token_type) -> tuple[str, bool, bool] | None:
    """Walk up the Pygments token hierarchy until we find a matching style,
    or return ``None`` if nothing in the ancestry has been styled."""
    while token_type is not None:
        if token_type in _TOKEN_STYLE:
            return _TOKEN_STYLE[token_type]
        token_type = token_type.parent
    return None


def _code_language(node: nodes.Element) -> str | None:
    """Return the lexer name for a ``literal_block`` node.

    RST's ``code-block::`` sets ``node['language']`` directly. MyST's fenced
    code blocks instead stash the language as a class alongside ``'code'``
    (e.g. ``classes=['code', 'python']``), so we fall back to the first
    non-``'code'`` class.
    """
    explicit = node.get('language')
    if explicit:
        return explicit
    for cls in node.get('classes', []):
        if cls != 'code':
            return cls
    return None


def _emit_highlighted(paragraph, source: str, language: str | None) -> None:
    """Tokenize ``source`` via the Pygments lexer for ``language`` and
    append coloured runs to ``paragraph``. Newlines within token text are
    converted to ``<w:br/>`` elements so the code block wraps correctly
    inside a single paragraph."""
    try:
        lexer = get_lexer_by_name(language, stripall=False) if language else TextLexer()
    except ClassNotFound:
        lexer = TextLexer()

    for token_type, text in lex(source, lexer):
        if not text:
            continue
        style = _token_style(token_type)
        parts = text.split('\n')
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                run.font.name = CODE_FONT_NAME
                if style is not None:
                    color_hex, bold, italic = style
                    run.font.color.rgb = RGBColor.from_string(color_hex)
                    if bold:
                        run.bold = True
                    if italic:
                        run.italic = True
            if i < len(parts) - 1:
                # Newline between token pieces → <w:br/>
                paragraph.add_run().add_break()


def _collect_styled_chunks(node: nodes.Node):
    """Walk an inline docutils subtree and yield ``(text, bold, italic, strike)``
    tuples for each ``Text`` descendant, inheriting ``bold`` / ``italic``
    from ``strong``/``emphasis``/``literal_*`` ancestors and tracking
    ``strike`` state across siblings (MyST emits strikethrough as raw-HTML
    ``<s>``/``</s>`` sibling markers rather than a wrapping node).

    Used to preserve inline formatting inside hyperlink text and footnote
    bodies, where the normal visitor-driven run emission can't run.
    """
    state = {'strike': False}
    yield from _walk_inline(node, bold=False, italic=False, state=state)


def _walk_inline(node, *, bold, italic, state):
    if isinstance(node, nodes.Text):
        yield (node.astext(), bold, italic, state['strike'])
        return
    # Use tagname rather than isinstance so Sphinx-only classes
    # (``sphinx.addnodes.literal_strong`` / ``literal_emphasis``) pick up
    # without forcing a Sphinx import on the writer module.
    tag = getattr(node, 'tagname', None)
    if tag in ('strong', 'literal_strong'):
        bold = True
    elif tag in ('emphasis', 'literal_emphasis'):
        italic = True
    elif isinstance(node, nodes.raw) and node.get('format') == 'html':
        html = node.astext().strip().lower()
        if html in ('<s>', '<strike>', '<del>'):
            state['strike'] = True
        elif html in ('</s>', '</strike>', '</del>'):
            state['strike'] = False
        return
    for child in node.children:
        yield from _walk_inline(child, bold=bold, italic=italic, state=state)


def _option_group_signature(option_group: nodes.Element) -> str:
    """Flatten an ``option_group`` to ``-o FILE, --out=FILE``-style text.

    Each ``option`` child contributes its ``option_string`` followed by an
    optional ``option_argument``. docutils already embeds the delimiter
    (``' '`` or ``'='``) at the start of ``option_argument.astext()``, so
    the string concatenates without extra separators. Multiple options in
    the group are joined by commas.
    """
    parts: list[str] = []
    for option in option_group.findall(nodes.option):
        parts.append(option.astext())
    return ', '.join(parts)


ADMONITION_LABELS = {
    'attention': 'Attention',
    'caution':   'Caution',
    'danger':    'Danger',
    'error':     'Error',
    'hint':      'Hint',
    'important': 'Important',
    'note':      'Note',
    'tip':       'Tip',
    'warning':   'Warning',
    'seealso':   'See also',
}


class DocxWriter(writers.Writer):
    """docutils writer producing a python-docx ``Document``."""

    supported = ('docx',)
    settings_spec = ('No options here.', '', ())
    settings_defaults: dict = {}

    output = None

    def __init__(self, builder):
        writers.Writer.__init__(self)
        self.builder = builder
        self.template_path: Path | None = self._resolve_template_path()

        if self.template_path is None:
            self.docx_container = Document()
        else:
            self.docx_container = Document(str(self.template_path))

    def _resolve_template_path(self) -> Path | None:
        """Resolve ``docx_template`` against Sphinx's standard template lookup.

        Search order for a relative ``docx_template``:

        1. Each entry of ``templates_path`` (resolved against ``srcdir``),
           in the order they appear in ``conf.py``.
        2. ``srcdir`` itself (for back-compat with pre-2.1 behaviour and
           projects that don't set ``templates_path``).

        The first file that exists wins. If none exist, we fall through
        to ``srcdir/<dotx>`` so python-docx raises a clear "not found"
        error pointing at the expected location rather than silently
        using a default template.
        """
        dotx = self.builder.config['docx_template']
        if not dotx:
            return None
        template = Path(dotx)
        if template.is_absolute():
            logger.info("using docx template: %s", template)
            return template

        srcdir = Path(self.builder.env.srcdir)
        templates_path = getattr(self.builder.config, 'templates_path', None) or []
        if isinstance(self.builder.config, dict):
            templates_path = self.builder.config.get('templates_path') or []

        search_dirs = [srcdir / entry for entry in templates_path]
        search_dirs.append(srcdir)
        for base in search_dirs:
            candidate = base / template
            if candidate.is_file():
                logger.info("using docx template: %s", candidate)
                return candidate

        fallback = srcdir / template
        logger.warning(
            "docx_template %r not found; searched %s",
            dotx, ', '.join(str(d) for d in search_dirs),
        )
        return fallback

    def save(self, filename):
        self.docx_container.save(filename)

    def translate(self):
        visitor = DocxTranslator(
                self.document, self.builder, self.docx_container)
        self.document.walkabout(visitor)
        self.output = ''  # visitor.body


class DocxState:
    """Tracks which part of the document paragraphs are being appended to.

    Used to support nested output contexts — e.g. a list inside a table cell.
    """
    def __init__(self, location=None):
        self.location = location
        self.table = None
        self.column_widths = None
        self.table_style = None
        self.more_cols = 0
        self.row = None
        self.cell_counter = 0
        self.ncolumns = 1
        "Number of columns in the current table."
        self.row_index = -1
        "Zero-based index of the row currently being visited (incremented in visit_row)."
        self.rowspan_until: dict[int, int] = {}
        "col_idx → last row_index still occupied by an earlier rowspan."
        self.pending_vmerges: list[dict] = []
        "Vertical-merge plan applied in depart_table once all rows exist."


# noinspection PyClassicStyleClass,PyMissingOrEmptyDocstring,PyUnusedLocal
class DocxTranslator(nodes.NodeVisitor):
    """Visitor class to create docx content."""

    def __init__(self, document, builder, docx_container):
        self.builder = builder
        self.docx_container = docx_container
        nodes.NodeVisitor.__init__(self, document)

        # TODO: Perhaps move the list_style into DocxState.
        # However, it should still be a list, and not a separate state,
        # because nested lists are not really nested.
        # So it will only be necessary if there are lists in tables
        # that are in lists.
        self.list_style = []
        self.list_level = 0

        # TODO: And what about sectionlevel?
        self.sectionlevel = 0

        self.table_style_default = 'Grid Table 4'
        self.in_literal_block = False
        self.strong = False
        self.emphasis = False
        self.subscript = False
        self.superscript = False
        self.strike = False

        self.current_state = DocxState(location=self.docx_container)
        self.current_state.table_style = self.table_style_default

        "The place where paragraphs will be added."
        self.old_states = []
        "A list of older states, e.g. typically [document, table-cell]"

        self.current_paragraph = None
        "The current paragraph that text is being added to."

        self._pending_text_prefix: str | None = None
        self._desc_param_count = 0  # parameter counter for a desc_parameterlist
        "Text to prepend on the next add_text call, then clear. Used by visit_list_item to emit a task-list checkbox glyph that lands in the same paragraph as the task text."

        self._bookmark_counter = 0
        "Monotonic OOXML bookmark id; incremented each time a bookmark is emitted."
        self._emitted_bookmarks: set[str] = set()
        "Names already bookmarked, to avoid duplicates when a section has multiple ids."

        self._figure_counter = 0
        "Monotonic figure number assigned by visit_figure; mirrors the value that Word's SEQ field will render."
        self._figure_numbers: dict[str, int] = {}
        "figure_id → figure_number mapping, for future :numref: resolution."

        self._footnote_counter = 0
        "Monotonic OOXML footnote id (ids 0 and -1 are reserved for separators)."
        self._footnote_ooxml_id: dict[str, int] = {}
        "docutils footnote/citation id → OOXML id (stable across reference and body visits)."
        self._footnotes_root = None
        "Lazy handle to <w:footnotes> in footnotes.xml; created on first use."

    def add_text(self, text):
        dprint()
        if self._pending_text_prefix is not None:
            text = self._pending_text_prefix + text
            self._pending_text_prefix = None
        textrun = self.current_paragraph.add_run(text)
        if self.strong:
            textrun.bold = True
        if self.emphasis:
            textrun.italic = True
        if self.subscript:
            textrun.font.subscript = True
        if self.superscript:
            textrun.font.superscript = True
        if self.strike:
            textrun.font.strike = True

    def new_state(self, location):
        dprint()
        self.old_states.append(self.current_state)
        self.current_state = DocxState(location=location)

    def end_state(self, first=None):
        dprint()
        self.current_state = self.old_states.pop()

    def visit_start_of_file(self, node):
        dprint()
        # TODO: HB should visit_start_of_file reset the sectionlevel?
        # If so, should it start a new state? If so, with which location?

        # FIXME: visit_start_of_file not close previous section.
        # sectionlevel keep previous and new file's heading level start with
        # previous + 1.
        # This quick hack reset sectionlevel per file.
        # (BTW Sphinx has heading levels per file? or entire document?)
        self.sectionlevel = 0

    def depart_start_of_file(self, node):
        dprint()

    def visit_document(self, node):
        dprint()

    def depart_document(self, node):
        dprint()

    def visit_highlightlang(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_section(self, node):
        dprint()
        self.sectionlevel += 1

    def depart_section(self, node):
        dprint()
        if self.sectionlevel > 0:
            self.sectionlevel -= 1

    def visit_topic(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_topic(self, node):
        dprint()
        raise nodes.SkipNode

    visit_sidebar = visit_topic
    depart_sidebar = depart_topic

    def visit_rubric(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('-[ ')

    def depart_rubric(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(' ]-')

    def visit_compound(self, node):
        dprint()
        pass

    def depart_compound(self, node):
        dprint()
        pass

    def visit_glossary(self, node):
        dprint()
        pass

    def depart_glossary(self, node):
        dprint()
        pass

    def visit_title(self, node):
        dprint()
        self.current_paragraph = self.current_state.location.add_heading(level=self.sectionlevel)
        # Drop a bookmark for each of the parent section's ids so internal
        # references (``:ref:``, ``[text](#id)``) can anchor at this heading.
        parent = node.parent
        if isinstance(parent, nodes.section):
            for anchor_id in parent.get('ids', []):
                self._emit_bookmark(self.current_paragraph, anchor_id)

    def depart_title(self, node):
        dprint()

    def _emit_bookmark(self, paragraph, name: str) -> None:
        """Emit a zero-width bookmark in ``paragraph`` if ``name`` hasn't
        been bookmarked yet in this document."""
        if name in self._emitted_bookmarks:
            return
        self._emitted_bookmarks.add(name)
        add_bookmark(paragraph, name, self._bookmark_counter)
        self._bookmark_counter += 1

    def visit_subtitle(self, node):
        dprint()
        pass

    def depart_subtitle(self, node):
        dprint()
        pass

    def visit_attribution(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('-- ')

    def depart_attribution(self, node):
        dprint()
        pass

    # -- autodoc / py-domain descriptions --------------------------------
    # Sphinx's `sphinx.ext.autodoc` and domain directives (`.. py:function::`
    # etc.) emit a `desc` container with a `desc_signature` header and a
    # `desc_content` body. We render the signature as a bold paragraph
    # (name + parameters + return type), hide the dotted module prefix
    # (`desc_addname`) for readability, and let the body's paragraphs +
    # field_list flow at the current location. Previously every desc_*
    # visitor was SkipNode, so autodoc output vanished entirely — see
    # GitHub issue #16.

    def visit_desc(self, node):
        dprint()

    def depart_desc(self, node):
        dprint()

    def visit_desc_signature(self, node):
        """Start a new paragraph for the signature line; children populate it."""
        dprint()
        self.current_paragraph = self.current_state.location.add_paragraph()
        self._desc_param_count = 0

    def depart_desc_signature(self, node):
        dprint()

    def visit_desc_name(self, node):
        """The canonical name (function / class / attribute) — bold."""
        dprint()
        self.current_paragraph.add_run(node.astext()).bold = True
        raise nodes.SkipNode

    def depart_desc_name(self, node):
        dprint()

    def visit_desc_addname(self, node):
        """The module-qualified prefix (e.g. ``foo.bar.``). Dropped for
        readability — the containing section usually provides the module."""
        dprint()
        raise nodes.SkipNode

    def depart_desc_addname(self, node):
        dprint()

    def visit_desc_type(self, node):
        """Inline type annotation on an attribute / variable description."""
        dprint()
        self.current_paragraph.add_run(node.astext())
        raise nodes.SkipNode

    def depart_desc_type(self, node):
        dprint()

    def visit_desc_returns(self, node):
        """Return-type annotation appended to the signature, e.g. ``-> bool``."""
        dprint()
        self.current_paragraph.add_run(f' → {node.astext()}')
        raise nodes.SkipNode

    def depart_desc_returns(self, node):
        dprint()

    def visit_desc_parameterlist(self, node):
        """Open the parameter list with ``(``; children add params + commas."""
        dprint()
        self.current_paragraph.add_run('(')
        self._desc_param_count = 0

    def depart_desc_parameterlist(self, node):
        dprint()
        self.current_paragraph.add_run(')')

    def visit_desc_parameter(self, node):
        """One parameter; inject a ``, `` separator if we've already emitted one."""
        dprint()
        if self._desc_param_count > 0:
            self.current_paragraph.add_run(', ')
        self._desc_param_count += 1
        self.current_paragraph.add_run(node.astext())
        raise nodes.SkipNode

    def visit_desc_optional(self, node):
        """Wrap optional parameters in ``[...]`` per Python signature conventions."""
        dprint()
        self.current_paragraph.add_run('[')

    def depart_desc_optional(self, node):
        dprint()
        self.current_paragraph.add_run(']')

    def visit_desc_annotation(self, node):
        """Inline annotation like ``class `` before a class signature."""
        dprint()
        self.current_paragraph.add_run(node.astext())
        raise nodes.SkipNode

    def depart_desc_annotation(self, node):
        dprint()

    def visit_refcount(self, node):
        dprint()

    def depart_refcount(self, node):
        dprint()

    def visit_desc_content(self, node):
        """Body of a description: docstring paragraphs + field list. Walked
        normally — paragraphs emit as regular paragraphs and field_list
        renders as a 2-col table (see visit_field_list)."""
        dprint()

    def depart_desc_content(self, node):
        dprint()

    def visit_figure(self, node):
        """Record a figure number for the `ids` on this node so :numref:
        resolution (handled by Sphinx's env before we see the tree) can find
        a matching bookmark, and so captions render a consistent number."""
        dprint()
        self._figure_counter += 1
        for fid in node.get('ids', []):
            self._figure_numbers[fid] = self._figure_counter

    def depart_figure(self, node):
        dprint()

    def visit_caption(self, node):
        """Emit the caption as a dedicated paragraph (``Caption`` style when
        the template provides one) prefixed by a bold ``Figure <N>. `` with
        Word's ``SEQ Figure`` field so the number auto-refreshes, and drop
        bookmarks for each of the parent figure's ids.
        """
        dprint()
        style = 'Caption'
        try:
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError:
            style = None

        para = self.current_state.location.add_paragraph(style=style)
        para.add_run('Figure ').bold = True
        add_seq_field(para, 'Figure', initial_value=self._figure_counter)
        para.add_run('. ').bold = True

        parent = node.parent
        if isinstance(parent, nodes.figure):
            for fid in parent.get('ids', []):
                self._emit_bookmark(para, fid)

        self.current_paragraph = para

    def depart_caption(self, node):
        dprint()

    def visit_productionlist(self, node):
        dprint()
        raise nodes.SkipNode
        # names = []
        # for production in node:
        #     names.append(production['tokenname'])
        # maxlen = max(len(name) for name in names)
        # for production in node:
        #     if production['tokenname']:
        #         self.add_text(production['tokenname'].ljust(maxlen) + ' ::=')
        #         lastname = production['tokenname']
        #     else:
        #         self.add_text('%s    ' % (' '*len(lastname)))
        #     self.add_text(production.astext() + '\n')
        # raise nodes.SkipNode

    def visit_seealso(self, node):
        dprint()

    def depart_seealso(self, node):
        dprint()

    def _allocate_footnote_id(self, docutils_id: str) -> int:
        """Get (or allocate) the OOXML footnote id for a given docutils id.
        Ensures the ``footnotes.xml`` part exists on first call."""
        if self._footnotes_root is None:
            self._footnotes_root = ensure_footnotes_part(self.docx_container)
        if docutils_id not in self._footnote_ooxml_id:
            self._footnote_counter += 1
            self._footnote_ooxml_id[docutils_id] = self._footnote_counter
        return self._footnote_ooxml_id[docutils_id]

    def _footnote_body_chunks(self, node):
        """Collect the footnote body as ``StyledChunk`` tuples, skipping the
        leading ``<label>`` child (Word auto-numbers) and inserting a space
        between top-level body paragraphs so adjacent paragraphs don't
        concatenate without whitespace."""
        chunks: list[tuple[str, bool, bool, bool]] = []
        for i, child in enumerate(node.children):
            if isinstance(child, nodes.label):
                continue
            if i > 0 and chunks:
                # Separator between paragraphs of the same footnote body.
                chunks.append((' ', False, False, False))
            for text, bold, italic, strike in _collect_styled_chunks(child):
                if text:
                    chunks.append((text, bold, italic, strike))
        return chunks

    def visit_footnote(self, node):
        """Emit a ``<w:footnote>`` entry into ``footnotes.xml`` and skip
        re-emitting the body in the main document."""
        dprint()
        ids = node.get('ids', [])
        if not ids:
            raise nodes.SkipNode
        ooxml_id = self._allocate_footnote_id(ids[0])
        chunks = self._footnote_body_chunks(node)
        add_footnote(self._footnotes_root, ooxml_id, chunks or '')
        raise nodes.SkipNode

    def depart_footnote(self, node):
        dprint()

    def visit_citation(self, node):
        """Citations share the footnote rendering path — Word doesn't
        distinguish them natively, and the docutils structure is the same."""
        dprint()
        ids = node.get('ids', [])
        if not ids:
            raise nodes.SkipNode
        ooxml_id = self._allocate_footnote_id(ids[0])
        chunks = self._footnote_body_chunks(node)
        add_footnote(self._footnotes_root, ooxml_id, chunks or '')
        raise nodes.SkipNode

    def depart_citation(self, node):
        dprint()

    def visit_label(self, node):
        dprint()
        raise nodes.SkipNode

    # XXX: option list could use some better styling

    def visit_option_list(self, node):
        """Render an option list as a 2-column table: signature | description."""
        dprint()
        self.new_state(location=self.current_state.location)
        self.current_state.table = self.current_state.location.add_table(rows=0, cols=2)

    def depart_option_list(self, node):
        dprint()
        self.end_state()

    def visit_option_list_item(self, node):
        dprint()
        self.current_state.row = self.current_state.table.add_row()

    def depart_option_list_item(self, node):
        dprint()

    def visit_option_group(self, node):
        """Flatten the option_group tree to a ``-o FILE, --out=FILE`` style
        signature and emit it into column 0. Skip children — we've rendered."""
        dprint()
        signature = _option_group_signature(node)
        cell = self.current_state.row.cells[0]
        cell.paragraphs[0].add_run(signature)
        raise nodes.SkipNode

    def depart_option_group(self, node):
        dprint()

    def visit_option(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_option(self, node):
        dprint()

    def visit_option_string(self, node):
        dprint()

    def depart_option_string(self, node):
        dprint()

    def visit_option_argument(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_option_argument(self, node):
        dprint()

    def visit_description(self, node):
        dprint()
        cell = self.current_state.row.cells[1]
        self.new_state(location=cell)
        if not cell.paragraphs[0].text:
            self.current_paragraph = cell.paragraphs[0]
        else:
            self.current_paragraph = cell.add_paragraph()

    def depart_description(self, node):
        dprint()
        self.end_state()

    def visit_tabular_col_spec(self, node):
        dprint()
        # TODO: properly implement this!!
        spec = node['spec']
        widths = [float(piece.split('cm')[0]) for piece in spec.split("{")[1:]]
        self.current_state.column_widths = widths
        raise nodes.SkipNode

    def visit_colspec(self, node):
        dprint()
        # The difficulty here is getting the right column width.
        # This can be specified with a tabular_col_spec, see above.
        #
        # Otherwise it is derived from the number of columns, which is
        # defined in visit_tgroup (a bit hackish).
        # The _block_width is the full width of the document, and this
        # is divided by the number of columns.
        #
        # It would perhaps also be possible to use node['colwidth'] in some way.
        # node['colwidth'] contains an integer like 22, the width of the column in ascii
        if self.current_state.column_widths:
            width = self.current_state.column_widths[0]
            self.current_state.column_widths = self.current_state.column_widths[1:]
            self.current_state.table.add_column(Cm(width))
        else:
            # noinspection PyProtectedMember
            self.current_state.table.add_column(
                self.docx_container._block_width // self.current_state.ncolumns
            )

        raise nodes.SkipNode

    def depart_colspec(self, node):
        dprint()

    def visit_tgroup(self, node):
        dprint()
        colspecs = [c for c in node.children if isinstance(c, nodes.colspec)]
        self.current_state.ncolumns = len(colspecs)

    def depart_tgroup(self, node):
        dprint()
        self.current_state.ncolumns = 1
        pass

    def visit_thead(self, node):
        dprint()
        pass

    def depart_thead(self, node):
        dprint()
        pass

    def visit_tbody(self, node):
        dprint()

    def depart_tbody(self, node):
        dprint()
        pass

    def visit_row(self, node):
        dprint()
        self.current_state.row = self.current_state.table.add_row()
        self.current_state.cell_counter = 0
        self.current_state.row_index += 1

    def depart_row(self, node):
        dprint()
        pass

    def visit_entry(self, node):
        dprint()
        morerows = node.get('morerows', 0)
        morecols = node.get('morecols', 0)
        if morecols:
            self.current_state.more_cols = morecols

        # Skip cell-counter positions that are still occupied by an earlier
        # row's rowspan. ``rowspan_until[col]`` holds the last row_index
        # still covered by that column's merge.
        while (
            self.current_state.rowspan_until.get(self.current_state.cell_counter, -1)
            >= self.current_state.row_index
        ):
            self.current_state.cell_counter += 1

        cell = self.current_state.row.cells[self.current_state.cell_counter]
        if self.current_state.more_cols:
            for i in range(self.current_state.more_cols):
                cell = cell.merge(
                    self.current_state.row.cells[self.current_state.cell_counter + i + 1]
                )

        if morerows:
            # Record the vertical merge plan (applied in depart_table once all
            # rows exist) and mark every column covered by this cell (plus
            # any horizontal span) as occupied for the next ``morerows`` rows.
            self.current_state.pending_vmerges.append({
                'top_row': self.current_state.row_index,
                'col': self.current_state.cell_counter,
                'row_span': morerows,
                'col_span': morecols,
            })
            last_row = self.current_state.row_index + morerows
            for delta in range(morecols + 1):
                col = self.current_state.cell_counter + delta
                self.current_state.rowspan_until[col] = last_row

        self.new_state(location=cell)
        # For some annoying reason, a new paragraph is automatically added
        # to each table cell. This is frustrating when you want, e.g. to
        # add a list item instead of a normal paragraph.
        self.current_paragraph = cell.paragraphs[0]

    def depart_entry(self, node):
        dprint()
        self.end_state()
        self.current_state.cell_counter = self.current_state.cell_counter + self.current_state.more_cols + 1
        self.current_state.more_cols = 0

    def visit_table(self, node):
        dprint()

        style = self.current_state.table_style
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.TABLE)
        except KeyError as exc:
            logger.warning('style "%s" is missing (%r); falling back to no style', style, exc)
            style = None

        # Columns are added when a colspec is visited.

        # It is only possible to use a style in add_table when adding a
        # table to the root document. That is, not for a table in a table.
        if len(self.old_states):
            self.current_state.table = self.current_state.location.add_table(rows=0, cols=0)
        else:
            self.current_state.table = self.current_state.location.add_table(
                rows=0, cols=0, style=style)

    def depart_table(self, node):
        dprint()

        # Apply any pending vertical merges now that every row exists.
        for plan in self.current_state.pending_vmerges:
            top_row = self.current_state.table.rows[plan['top_row']]
            top_cell = top_row.cells[plan['col']]
            for i in range(1, plan['row_span'] + 1):
                target_row = self.current_state.table.rows[plan['top_row'] + i]
                for j in range(plan['col_span'] + 1):
                    target_cell = target_row.cells[plan['col'] + j]
                    top_cell.merge(target_cell)

        self.current_state.table = None
        self.current_state.table_style = self.table_style_default
        self.current_state.rowspan_until = {}
        self.current_state.pending_vmerges = []
        self.current_state.row_index = -1

        # Add an empty paragraph to prevent tables from being concatenated.
        # TODO: Figure out some better solution.
        self.current_state.location.add_paragraph("")

    def visit_acks(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text(', '.join(n.astext() for n in node.children[0].children)
        #               + '.')

    def visit_image(self, node):
        dprint()
        uri = node.attributes['uri']
        file_path = Path(self.builder.env.srcdir) / uri

        width = _length_to_emu(node.get('width'))
        height = _length_to_emu(node.get('height'))

        # :scale: and % widths require the image's intrinsic dimensions to
        # compute a target size. python-docx doesn't expose that without
        # actually opening the image; we skip these for now and fall back
        # to native sizing. Surface a single log line so users don't think
        # the attribute was silently honoured.
        raw_width = node.get('width', '')
        if raw_width.endswith('%'):
            logger.warning(
                'image %s: percentage width %r ignored (not supported); '
                'using native size', uri, raw_width,
            )
        if node.get('scale') and not (width or height):
            logger.warning(
                'image %s: :scale: %r ignored (not supported); '
                'using native size', uri, node.get('scale'),
            )

        # Inline image inside a paragraph (e.g. ``![alt](x) mid-sentence``)
        # must go into the running paragraph as a run; block images get
        # their own paragraph at the current location.
        inline = (
            isinstance(node.parent, nodes.paragraph)
            and self.current_paragraph is not None
        )
        para = (
            self.current_paragraph
            if inline
            else self.current_state.location.add_paragraph()
        )

        run = para.add_run()
        try:
            shape = run.add_picture(str(file_path), width=width, height=height)
        except (OSError, ValueError) as exc:
            logger.warning('failed to embed image %s: %s', file_path, exc)
            return

        # :align: center|left|right — only meaningful on block-level images.
        if not inline:
            align = node.get('align')
            if align in _ALIGN_MAP:
                para.alignment = _ALIGN_MAP[align]

        # Alt text / title text — set on the inline shape's wp:docPr element.
        alt = node.get('alt')
        if alt:
            try:
                doc_pr = shape._inline.docPr  # noqa: SLF001
                doc_pr.set('descr', alt)
                doc_pr.set('title', alt)
            except AttributeError:
                pass  # python-docx version without the expected attribute

    def depart_image(self, node):
        dprint()

    def visit_transition(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('=' * 70)

    def visit_bullet_list(self, node):
        dprint()
        self.list_style.append('bullet')
        self.list_level += 1

    def depart_bullet_list(self, node):
        dprint()
        self.list_style.pop()
        self.list_level -= 1

    def visit_enumerated_list(self, node):
        dprint()
        self.list_style.append('number')
        self.list_level += 1

    def depart_enumerated_list(self, node):
        dprint()
        self.list_style.pop()
        self.list_level -= 1

    def visit_definition_list(self, node):
        """Render a definition list as a 2-column table (term | definition)."""
        dprint()
        # Push a fresh state so nested deflists don't clobber our table handle.
        self.new_state(location=self.current_state.location)
        self.current_state.table = self.current_state.location.add_table(rows=0, cols=2)

    def depart_definition_list(self, node):
        dprint()
        self.end_state()

    def visit_list_item(self, node):
        dprint()
        # A new paragraph is created here, but the next visit is to
        # paragraph, so that would add another paragraph. That is
        # prevented if current_paragraph is an empty List paragraph.
        # Pick the Word list style based on the innermost list's kind.
        # ``self.list_style`` is a stack maintained by visit_bullet_list /
        # visit_enumerated_list. Bare ``'List Number'`` / ``'List Bullet'``
        # render at depth 1; deeper nesting uses the ``N`` suffix Word's
        # built-in styles define.
        base = 'List Number' if self.list_style and self.list_style[-1] == 'number' else 'List Bullet'
        style = base if self.list_level < 2 else f'{base} {self.list_level}'
        try:
            # Check whether the style is part of the document.
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            logger.warning('style "%s" is missing (%r); falling back to no style', style, exc)
            style = None

        curloc = self.current_state.location
        if isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    self.current_paragraph = curloc.paragraphs[0]
                    self.current_paragraph.style = style
                else:
                    self.current_paragraph = curloc.add_paragraph(style=style)
            else:
                self.current_paragraph = curloc.add_paragraph(style=style)
        else:
            self.current_paragraph = curloc.add_paragraph(style=style)

        # GFM task-list (MyST `tasklist` extension): list_item carries the
        # "task-list-item" class and its first paragraph starts with a raw
        # <input type="checkbox" [checked]> node (which visit_raw drops).
        # Queue a Unicode ballot-box glyph for the next add_text call so the
        # glyph + task text end up in the same paragraph.
        if 'task-list-item' in node.get('classes', []):
            self._pending_text_prefix = _checkbox_prefix(node)

    def depart_list_item(self, node):
        dprint()

    def visit_definition_list_item(self, node):
        """A new ``<term>/<definition>`` pair — add a row to the deflist table."""
        dprint()
        self.current_state.row = self.current_state.table.add_row()

    def depart_definition_list_item(self, node):
        dprint()

    def visit_term(self, node):
        """Term text goes into column 0 as a bold run; any inline formatting
        inside the term is flattened to plain text for simplicity."""
        dprint()
        cell = self.current_state.row.cells[0]
        para = cell.paragraphs[0]
        para.add_run(node.astext()).bold = True
        raise nodes.SkipNode

    def depart_term(self, node):
        dprint()

    def visit_classifier(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_classifier(self, node):
        dprint()

    def visit_definition(self, node):
        """Definition body flows into column 1. Multiple ``<definition>``
        children (MyST emits one per colon-marked line) land as consecutive
        paragraphs in the same cell."""
        dprint()
        cell = self.current_state.row.cells[1]
        self.new_state(location=cell)
        if not cell.paragraphs[0].text:
            self.current_paragraph = cell.paragraphs[0]
        else:
            self.current_paragraph = cell.add_paragraph()

    def depart_definition(self, node):
        dprint()
        self.end_state()

    # -- field lists -----------------------------------------------------
    # Render field_list as a 2-col table (name | body). Primary consumer
    # is autodoc's ``:param:``/``:returns:``/``:raises:`` fields, but
    # docutils document metadata (``:author:``/``:version:``) and any
    # user-provided field list render the same way. The pattern mirrors
    # visit_definition_list.

    def visit_field_list(self, node):
        dprint()
        self.new_state(location=self.current_state.location)
        self.current_state.table = self.current_state.location.add_table(rows=0, cols=2)

    def depart_field_list(self, node):
        dprint()
        self.end_state()

    def visit_field(self, node):
        dprint()
        self.current_state.row = self.current_state.table.add_row()

    def depart_field(self, node):
        dprint()

    def visit_field_name(self, node):
        dprint()
        cell = self.current_state.row.cells[0]
        cell.paragraphs[0].add_run(node.astext()).bold = True
        raise nodes.SkipNode

    def depart_field_name(self, node):
        dprint()

    def visit_field_body(self, node):
        dprint()
        cell = self.current_state.row.cells[1]
        self.new_state(location=cell)
        if not cell.paragraphs[0].text:
            self.current_paragraph = cell.paragraphs[0]
        else:
            self.current_paragraph = cell.add_paragraph()

    def depart_field_body(self, node):
        dprint()
        self.end_state()

    def visit_centered(self, node):
        dprint()
        pass

    def depart_centered(self, node):
        dprint()
        pass

    def visit_hlist(self, node):
        dprint()
        pass

    def depart_hlist(self, node):
        dprint()
        pass

    def visit_hlistcol(self, node):
        dprint()
        pass

    def depart_hlistcol(self, node):
        dprint()
        pass

    def _visit_admonition(self, node):
        """Render a note/warning/tip/etc. as a 1×1 table containing a bold
        title paragraph followed by the admonition body content."""
        dprint()
        # Create the wrapper table at the current location (document or cell).
        table = self.current_state.location.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]

        # Determine the label:
        #  - specialised admonition (note/warning/tip/…) → map tagname to a
        #    canonical label from ADMONITION_LABELS.
        #  - generic `.. admonition:: Custom Heading` → pull the user's text
        #    out of the first-child title and remove it so visit_title doesn't
        #    render it a second time.
        if node.tagname == 'admonition':
            label = 'Note'
            if node.children and isinstance(node.children[0], nodes.title):
                label = node.children[0].astext()
                node.remove(node.children[0])
        else:
            label = ADMONITION_LABELS.get(node.tagname, node.tagname.title())

        # Enter the cell as the current location, then emit the bold title
        # into the auto-created empty paragraph that python-docx adds to any
        # new cell.
        self.new_state(location=cell)
        title_para = cell.paragraphs[0]
        title_para.add_run(label).bold = True
        self.current_paragraph = title_para

    def _depart_admonition(self, node):
        dprint()
        self.end_state()

    visit_admonition = _visit_admonition
    depart_admonition = _depart_admonition
    visit_attention = _visit_admonition
    depart_attention = _depart_admonition
    visit_caution = _visit_admonition
    depart_caution = _depart_admonition
    visit_danger = _visit_admonition
    depart_danger = _depart_admonition
    visit_error = _visit_admonition
    depart_error = _depart_admonition
    visit_hint = _visit_admonition
    depart_hint = _depart_admonition
    visit_important = _visit_admonition
    depart_important = _depart_admonition
    visit_note = _visit_admonition
    depart_note = _depart_admonition
    visit_tip = _visit_admonition
    depart_tip = _depart_admonition
    visit_warning = _visit_admonition
    depart_warning = _depart_admonition

    def visit_versionmodified(self, node):
        dprint()
        raise nodes.SkipNode
        # from sphinx.locale import admonitionlabels, versionlabels, _
        # if node.children:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + ': ')
        # else:
        #     self.add_text(
        #             versionlabels[node['type']] % node['version'] + '.')

    def depart_versionmodified(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_literal_block(self, node):
        """Render a fenced code block as a single paragraph with one run per
        Pygments token (colour + bold/italic driven by ``_TOKEN_STYLE``) and
        ``<w:br/>`` between source lines. The ``Preformatted Text`` paragraph
        style is applied when available so overall block appearance (margins,
        background, default font) follows the user's template."""
        dprint()
        self.in_literal_block = True

        style = 'Preformatted Text'
        try:
            self.docx_container.styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
        except KeyError as exc:
            logger.warning('style "%s" is missing (%r); falling back to no style', style, exc)
            style = None

        self.current_paragraph = self.current_state.location.add_paragraph(style=style)
        self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Tokenize + emit coloured runs. SkipNode prevents visit_Text from
        # appending the raw source again.
        _emit_highlighted(self.current_paragraph, node.astext(), _code_language(node))
        raise nodes.SkipNode

    def depart_literal_block(self, node):
        dprint()
        self.in_literal_block = False

    def visit_doctest_block(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_doctest_block(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_line_block(self, node):
        dprint()
        raise nodes.SkipNode

    def depart_line_block(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_line(self, node):
        dprint()
        pass

    def depart_line(self, node):
        dprint()
        pass

    def visit_block_quote(self, node):
        dprint()

    def depart_block_quote(self, node):
        dprint()

    def visit_compact_paragraph(self, node):
        dprint()

    def depart_compact_paragraph(self, node):
        dprint()

    def visit_paragraph(self, node):
        dprint()

        curloc = self.current_state.location

        if (
            self.current_paragraph is not None
            and 'List' in self.current_paragraph.style.name
            and not self.current_paragraph.text
        ):
            # This is the first paragraph in a list item, so do not create another one.
            pass
        elif isinstance(curloc, _Cell):
            if len(curloc.paragraphs) == 1:
                if not curloc.paragraphs[0].text:
                    # An empty paragraph is created when a Cell is created.
                    # Reuse this paragraph.
                    self.current_paragraph = curloc.paragraphs[0]
                else:
                    self.current_paragraph = curloc.add_paragraph()
            else:
                self.current_paragraph = curloc.add_paragraph()
            # HACK because the style is messed up, TODO FIX
            self.current_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self.current_paragraph.paragraph_format.left_indent = 0
        else:
            self.current_paragraph = curloc.add_paragraph()

    def depart_paragraph(self, node):
        dprint()

    def visit_target(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_index(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_substitution_definition(self, node):
        dprint()
        raise nodes.SkipNode

    def visit_pending_xref(self, node):
        dprint()
        pass

    def depart_pending_xref(self, node):
        dprint()
        pass

    def visit_reference(self, node):
        """References render as clickable ``w:hyperlink`` runs.

        External URL → ``w:hyperlink`` with ``r:id`` pointing at an external
        relationship. Internal ref (``refid``, or relative ``refuri``
        starting with ``#``) → ``w:hyperlink`` with ``w:anchor`` pointing
        at a bookmark emitted from ``visit_title`` on the matching section.
        """
        dprint()
        refuri = node.get('refuri', '')
        refid = node.get('refid', '')

        if refuri and _EXTERNAL_URL_RE.match(refuri):
            if self.current_paragraph is None:
                self.current_paragraph = self.current_state.location.add_paragraph()
            chunks = [c for c in _collect_styled_chunks(node) if c[0]]
            add_hyperlink(self.current_paragraph, refuri, chunks or node.astext())
            raise nodes.SkipNode

        # Internal anchor: take refid directly, or pull the fragment off an
        # intra-doc refuri like ``other.html#my-section`` / ``#my-section``.
        anchor = refid
        if not anchor and refuri and refuri.startswith('#'):
            anchor = refuri.lstrip('#')
        if anchor:
            if self.current_paragraph is None:
                self.current_paragraph = self.current_state.location.add_paragraph()
            chunks = [c for c in _collect_styled_chunks(node) if c[0]]
            add_internal_hyperlink(
                self.current_paragraph, anchor, chunks or node.astext()
            )
            raise nodes.SkipNode

    def depart_reference(self, node):
        dprint()

    # ``:numref:`` resolves to a ``nodes.number_reference`` — shape-compatible
    # with ``nodes.reference`` (has ``refid`` / ``refuri`` and a child
    # ``inline`` carrying the formatted text like "Fig. 1" / "Table 3").
    # Route it through the same hyperlink emitter so the Word reader clicks
    # back to the section bookmark at the figure id.
    visit_number_reference = visit_reference
    depart_number_reference = depart_reference

    def visit_download_reference(self, node):
        dprint()
        pass

    def depart_download_reference(self, node):
        dprint()
        pass

    def visit_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = True

    def depart_emphasis(self, node):
        dprint()
        # self.add_text('*')
        self.emphasis = False

    def visit_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def depart_literal_emphasis(self, node):
        dprint()
        # self.add_text('*')

    def visit_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = True

    def depart_strong(self, node):
        dprint()
        # self.add_text('**')
        self.strong = False

    def visit_abbreviation(self, node):
        dprint()
        # self.add_text('')

    def depart_abbreviation(self, node):
        dprint()
        # if node.hasattr('explanation'):
        #     self.add_text(' (%s)' % node['explanation'])

    def visit_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def depart_title_reference(self, node):
        dprint()
        # self.add_text('*')

    def visit_literal(self, node):
        dprint()
        # self.add_text('``')

    def depart_literal(self, node):
        dprint()
        # self.add_text('``')

    def visit_subscript(self, node):
        dprint()
        self.subscript = True

    def depart_subscript(self, node):
        dprint()
        self.subscript = False

    def visit_superscript(self, node):
        dprint()
        self.superscript = True

    def depart_superscript(self, node):
        dprint()
        self.superscript = False

    def visit_footnote_reference(self, node):
        """Insert a ``<w:footnoteReference w:id="…"/>`` run in the current
        paragraph. The OOXML id is allocated lazily and keyed by the target
        footnote's docutils id, so the ``visit_footnote`` body emission
        below uses the same id."""
        dprint()
        refid = node.get('refid')
        if not refid:
            raise nodes.SkipNode
        ooxml_id = self._allocate_footnote_id(refid)
        if self.current_paragraph is None:
            self.current_paragraph = self.current_state.location.add_paragraph()
        add_footnote_reference(self.current_paragraph, ooxml_id)
        raise nodes.SkipNode

    def visit_citation_reference(self, node):
        dprint()
        refid = node.get('refid')
        if not refid:
            raise nodes.SkipNode
        ooxml_id = self._allocate_footnote_id(refid)
        if self.current_paragraph is None:
            self.current_paragraph = self.current_state.location.add_paragraph()
        add_footnote_reference(self.current_paragraph, ooxml_id)
        raise nodes.SkipNode

    def visit_Text(self, node):
        dprint()
        text = node.astext()
        if not self.in_literal_block:
            # assert '\n\n' not in text, 'Found \n\n'
            # Replace double enter with single enter, and single enter with space.
            string_magic = 'TWOENTERSMAGICSTRING'
            text = text.replace('\n\n', string_magic)
            text = text.replace('\n', ' ')
            text = text.replace(string_magic, '\n')
        self.add_text(text)

    def depart_Text(self, node):
        dprint()
        pass

    def visit_generated(self, node):
        dprint()
        pass

    def depart_generated(self, node):
        dprint()
        pass

    def visit_inline(self, node):
        dprint()
        pass

    def depart_inline(self, node):
        dprint()
        pass

    def visit_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('>>')

    def depart_problematic(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<<')

    def visit_system_message(self, node):
        dprint()
        raise nodes.SkipNode
        # self.add_text('<SYSTEM MESSAGE: %s>' % node.astext())

    def visit_comment(self, node):
        """Consume RST ``.. comment`` nodes.

        A long-standing hack lets a comment carry ``DocxTableStyle <name>`` to
        override the per-document table style. Empty ``.. \\n`` comments
        previously raised ``IndexError`` when indexing ``node[0]`` — see GitHub
        issues #33 and #55. Use ``astext()`` which tolerates child-less
        comments.
        """
        dprint()
        text = node.astext()
        if 'DocxTableStyle' in text:
            self.current_state.table_style = text.split('DocxTableStyle')[-1].strip()
        raise nodes.SkipNode

    def visit_meta(self, node):
        dprint()
        raise nodes.SkipNode
        # only valid for HTML

    # -- math -------------------------------------------------------------
    # Renders inline ``:math:`` and block ``.. math::`` / MyST ``$…$`` /
    # ``$$…$$`` via pandoc's LaTeX → OMML conversion. The OMML element is
    # inserted directly into the current paragraph (inline) or a fresh
    # paragraph (display). When pandoc is unavailable or the LaTeX fails
    # to convert, the raw LaTeX source renders as a monospace fallback so
    # no content is silently dropped.

    def _render_math_fallback(self, latex: str, *, inline: bool) -> None:
        """Emit LaTeX source as a monospace run (or new paragraph for display)."""
        if inline:
            if self.current_paragraph is None:
                self.current_paragraph = self.current_state.location.add_paragraph()
            run = self.current_paragraph.add_run(latex)
        else:
            para = self.current_state.location.add_paragraph()
            run = para.add_run(latex)
        run.font.name = CODE_FONT_NAME
        logger.warning(
            'math rendered as plain LaTeX (pandoc unavailable or conversion failed): %r',
            latex[:80],
        )

    def visit_math(self, node):
        dprint()
        latex = node.astext()
        omml = latex_to_omml(latex, display=False)
        if omml is None:
            self._render_math_fallback(latex, inline=True)
            raise nodes.SkipNode
        if self.current_paragraph is None:
            self.current_paragraph = self.current_state.location.add_paragraph()
        self.current_paragraph._p.append(omml)
        raise nodes.SkipNode

    def depart_math(self, node):
        dprint()

    def visit_math_block(self, node):
        dprint()
        latex = node.astext()
        omml = latex_to_omml(latex, display=True)
        if omml is None:
            self._render_math_fallback(latex, inline=False)
            raise nodes.SkipNode
        para = self.current_state.location.add_paragraph()
        para._p.append(omml)
        raise nodes.SkipNode

    def depart_math_block(self, node):
        dprint()

    def visit_raw(self, node):
        """Drop raw content — except toggle the strike flag on HTML
        ``<s>`` / ``<strike>`` / ``<del>`` markers. MyST emits those
        around the text of ``~~strikethrough~~`` after logging a warning
        that strikethrough only renders natively in HTML; we treat them
        as a signal to the writer and render the text as struck runs."""
        dprint()
        if node.get('format') == 'html':
            html = node.astext().strip().lower()
            if html in ('<s>', '<strike>', '<del>'):
                self.strike = True
            elif html in ('</s>', '</strike>', '</del>'):
                self.strike = False
        raise nodes.SkipNode
        #     self.body.append(node.astext())

    def unknown_visit(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)

    def unknown_departure(self, node):
        dprint()
        raise nodes.SkipNode
        # raise NotImplementedError('Unknown node: ' + node.__class__.__name__)
